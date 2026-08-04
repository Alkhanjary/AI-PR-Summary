import io
import json
import re
import shlex
import shutil
import socket
import subprocess
import sys
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import urljoin, urlparse

import requests

from flask import Flask, jsonify, request, Response
from flask_cors import CORS

sys.path.insert(0, str(Path(__file__).resolve().parent))
from summarize_pr import (
    filter_diff,
    smart_truncate,
    scan_security,
    format_security_report,
    call_llm,
    SYSTEM_PROMPT,
    VULN_REPORT_SYSTEM_PROMPT,
    PROJECT_DETECT_SYSTEM_PROMPT,
    PROJECT_RETRY_SYSTEM_PROMPT,
    VULN_DETAIL_SYSTEM_PROMPT,
    AI_CODE_SCAN_SYSTEM_PROMPT,
    AI_WEB_SCAN_SYSTEM_PROMPT,
    AI_NETWORK_SCAN_SYSTEM_PROMPT,
    MAX_CHARS,
)
from openai import OpenAI
import os
from dotenv import load_dotenv

load_dotenv(Path(__file__).resolve().parent / ".env")

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', os.urandom(32).hex())

# Security: CSRF tokens, rate limiting, API key validation
import secrets
import hashlib
import time
from functools import wraps
from flask import session, abort

_csrf_tokens = {}  # token -> (creation_time, valid)
_request_counts = {}  # ip -> [(timestamp, endpoint)]
MAX_REQUESTS_PER_MINUTE = 60
CSRF_TOKEN_LIFETIME = 3600  # 1 hour

def _get_csrf_token():
  """Generate a new CSRF token."""
  token = secrets.token_urlsafe(32)
  _csrf_tokens[token] = (time.time(), True)
  return token

def _validate_csrf_token(token):
  """Validate a CSRF token."""
  if token not in _csrf_tokens:
    return False
  created, valid = _csrf_tokens[token]
  if not valid or (time.time() - created) > CSRF_TOKEN_LIFETIME:
    _csrf_tokens[token] = (created, False)
    return False
  return True

def require_csrf_token(f):
  """Decorator to require valid CSRF token in POST requests."""
  @wraps(f)
  def decorated_function(*args, **kwargs):
    if request.method == 'POST':
      token = request.form.get('csrf_token') or request.json.get('csrf_token')
      if not token or not _validate_csrf_token(token):
        abort(403)
    return f(*args, **kwargs)
  return decorated_function

def require_api_key(f):
  """Decorator to require valid API key for sensitive endpoints."""
  @wraps(f)
  def decorated_function(*args, **kwargs):
    key = request.headers.get('X-API-Key')
    expected = os.environ.get('API_KEY')
    if expected and key != expected:
      abort(401)
    return f(*args, **kwargs)
  return decorated_function

def rate_limit(endpoint_name):
  """Decorator to rate-limit requests (60 per minute per IP)."""
  def decorator(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
      ip = request.remote_addr
      now = time.time()
      if ip not in _request_counts:
        _request_counts[ip] = []
      # Clean old entries
      _request_counts[ip] = [ts for ts in _request_counts[ip] if now - ts < 60]
      if len(_request_counts[ip]) >= MAX_REQUESTS_PER_MINUTE:
        abort(429)  # Too many requests
      _request_counts[ip].append(now)
      return f(*args, **kwargs)
    return decorated_function
  return decorator

# This API can read arbitrary local folders and run build commands, so a
# wildcard CORS policy is genuinely dangerous: any website you happened to have
# open could call http://127.0.0.1:5000 and read the response. Restrict it to
# origins the dashboard is actually served from.
#
# "null" covers opening dashboard/index.html straight off disk (file:// sends
# Origin: null). That is how the README tells you to run it, so it has to work -
# but it is the weakest entry here, since a sandboxed iframe can also present a
# null origin. Serve the dashboard over http://127.0.0.1 and set
# DASHBOARD_ALLOW_FILE_ORIGIN=0 to drop it.
_allowed_origins = [
    re.compile(r"^https?://(localhost|127\.0\.0\.1)(:\d+)?$"),
]
if os.environ.get("DASHBOARD_ALLOW_FILE_ORIGIN", "1") != "0":
    _allowed_origins.append("null")
_extra_origins = [o.strip() for o in os.environ.get("DASHBOARD_ALLOWED_ORIGINS", "").split(",") if o.strip()]
CORS(app, origins=_allowed_origins + _extra_origins, supports_credentials=False)


@app.after_request
def _security_headers(resp):
    """Baseline hardening for the dashboard's own responses.

    The CSP matters most here: the dashboard is a single HTML file that renders
    data pulled from GitHub (PR titles, commit messages, branch names), so it
    limits the blast radius if any escaping is ever missed.

    Deliberately NOT set:
      - Strict-Transport-Security: this serves plain HTTP on loopback, and HSTS
        is ignored over HTTP anyway. Setting it would only risk poisoning the
        browser's HSTS state for localhost across every other project.
      - X-XSS-Protection: deprecated, removed from modern browsers, and its
        legacy auditor introduced vulnerabilities of its own. CSP replaces it.
    """
    resp.headers.setdefault("X-Content-Type-Options", "nosniff")
    resp.headers.setdefault("X-Frame-Options", "DENY")
    resp.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
    resp.headers.setdefault("Permissions-Policy",
                            "camera=(), microphone=(), geolocation=(), interest-cohort=()")
    resp.headers.setdefault(
        "Content-Security-Policy",
        "default-src 'self'; "
        # The dashboard is one self-contained file with inline <script>/<style>.
        "script-src 'self' 'unsafe-inline'; "
        "style-src 'self' 'unsafe-inline'; "
        "img-src 'self' data: https://avatars.githubusercontent.com; "
        "connect-src 'self' http://127.0.0.1:5000 http://localhost:5000; "
        "object-src 'none'; base-uri 'none'; frame-ancestors 'none'",
    )
    # Don't advertise the exact Werkzeug/Python versions.
    resp.headers["Server"] = "dashboard"
    return resp

# The exact interpreter running this server - used instead of a hardcoded
# "py" (Windows-launcher-only, not guaranteed even there) or "python3"
# (missing on some Windows setups) so scanner/detected-project subprocesses
# always resolve to a real, working Python on any platform.
PYTHON_CMD = sys.executable


def get_client():
    api_key = os.environ.get("LLM_API_KEY")
    base_url = os.environ.get("LLM_BASE_URL") or None
    if not api_key:
        return None, None
    model = os.environ.get("LLM_MODEL", "gpt-oss:20b-cloud")
    return OpenAI(api_key=api_key, base_url=base_url), model


@app.route("/api/health")
def health():
    return jsonify({"status": "ok"})


@app.route("/api/csrf-token", methods=["GET"])
def csrf_token():
    """Get a CSRF token for protected requests."""
    return jsonify({"csrf_token": _get_csrf_token()})


@app.route("/api/github/repos")
def github_repos():
    username = request.args.get("username", "")
    if not username:
        return jsonify({"error": "username required"}), 400
    r = requests.get(
        "https://api.github.com/users/" + username + "/repos",
        params={"per_page": 100, "sort": "updated", "direction": "desc"},
        headers=gh_headers(),
    )
    return jsonify(r.json()), r.status_code


@app.route("/api/github/user")
def github_user():
    """Powers the Overview tab's user profile card (name, bio, follower
    counts, account age, etc.) - separate from /api/github/repos, which
    only returns the repo list used to compute aggregate stats."""
    username = request.args.get("username", "")
    if not username:
        return jsonify({"error": "username required"}), 400
    r = requests.get(
        "https://api.github.com/users/" + username,
        headers=gh_headers(),
    )
    return jsonify(r.json()), r.status_code


@app.route("/api/github/user-activity")
def github_user_activity():
    """Powers the Overview tab's 'Involved in' section: pull requests the
    user has authored across ANY repo (not just ones they own), so
    contributions to other people's/orgs' repos show up too."""
    username = request.args.get("username", "")
    if not username:
        return jsonify({"error": "username required"}), 400
    r = requests.get(
        "https://api.github.com/search/issues",
        params={"q": "author:" + username + " is:pr", "sort": "updated", "order": "desc", "per_page": 20},
        headers=gh_headers(),
    )
    if not r.ok:
        return jsonify({"items": []}), r.status_code
    data = r.json()
    items = []
    for item in data.get("items", []):
        repo_url = item.get("repository_url", "")
        repo_full_name = "/".join(repo_url.split("/")[-2:]) if repo_url else ""
        items.append({
            "repo": repo_full_name,
            "title": item.get("title"),
            "html_url": item.get("html_url"),
            "state": "merged" if item.get("pull_request", {}).get("merged_at") else item.get("state"),
            "updated_at": item.get("updated_at"),
        })
    return jsonify({"items": items})


@app.route("/api/github/search-users")
def github_search_users():
    """Powers the username autocomplete dropdown: proxies GitHub's user
    search so the frontend doesn't need its own token/CORS handling."""
    q = request.args.get("q", "").strip()
    if len(q) < 2:
        return jsonify({"items": []})
    r = requests.get(
        "https://api.github.com/search/users",
        params={"q": q, "per_page": 6},
        headers=gh_headers(),
    )
    if not r.ok:
        return jsonify({"items": []}), r.status_code
    data = r.json()
    items = [
        {"login": item["login"], "avatar_url": item.get("avatar_url", "")}
        for item in data.get("items", [])
    ]
    return jsonify({"items": items})


def get_repo_dir(repo, branch=None):
    """Clones a fresh temporary working copy of the given owner/repo for
    this one operation. No on-disk cache is kept between calls (each call
    clones from scratch into a new temp directory) - this is what lets the
    Run scan tab follow whichever repo is picked up top, not just the
    AI-PR-Summary folder this server lives in. If branch is given, checks
    that branch out (defaults to whatever the clone's default branch is
    otherwise)."""
    if not repo:
        return Path(__file__).resolve().parent
    local_path = Path(tempfile.mkdtemp(prefix="repo_"))
    proc = subprocess.run(
        ["git", "clone", "--no-single-branch", f"https://github.com/{repo}.git", str(local_path)],
        capture_output=True, text=True,
    )
    # A failed clone used to pass silently, handing back an empty directory -
    # so the scan "succeeded" against nothing and reported no findings, which
    # looks identical to a clean repo. Fail loudly instead.
    if proc.returncode != 0:
        detail = (proc.stderr or proc.stdout or "").strip().splitlines()
        hint = detail[-1] if detail else f"git exited with code {proc.returncode}"
        raise RuntimeError(
            f"Could not clone https://github.com/{repo}.git - {hint}. "
            "Check the repo name, that it's public (or set GITHUB_TOKEN for private repos), "
            "and that git is installed and has network access."
        )
    if not any(local_path.iterdir()):
        raise RuntimeError(f"Clone of {repo} produced an empty folder - nothing to scan.")

    if branch:
        short = branch.split("/", 1)[1] if branch.startswith("origin/") else branch
        co = subprocess.run(["git", "checkout", short], capture_output=True, text=True, cwd=local_path)
        if co.returncode != 0:
            err = (co.stderr or "").strip().splitlines()
            raise RuntimeError(
                f"Cloned {repo}, but could not check out branch '{short}': "
                + (err[-1] if err else "unknown error")
            )
        subprocess.run(["git", "reset", "--hard", f"origin/{short}"], capture_output=True, text=True, cwd=local_path)
    return local_path


@app.route("/api/team-stats")
def team_stats():
    """Aggregates commit activity per contributor (commits, lines added,
    lines deleted) for the given repo, using GitHub's API for the commit
    list and per-commit stats. Powers the Monitor tab's team overview.
    Optional ?since=<ISO8601>&until=<ISO8601> filters to commits in that
    window (day/week/month/all-time presets, or a custom range, are all
    just different since/until values from the dashboard's point of view).
    Optional ?branch=<name> follows that branch instead of the repo's
    default branch."""
    repo = request.args.get("repo", "").strip()
    limit = int(request.args.get("limit", 100))
    since = request.args.get("since", "").strip()
    branch = request.args.get("branch", "").strip()
    if not repo:
        return jsonify({"error": "repo required"}), 400

    until = request.args.get("until", "").strip()
    params = {"per_page": min(limit, 100)}
    if since:
        params["since"] = since
    if until:
        params["until"] = until
    if branch:
        params["sha"] = branch.split("/", 1)[1] if branch.startswith("origin/") else branch

    r = requests.get(
        "https://api.github.com/repos/" + repo + "/commits",
        params=params,
        headers=gh_headers(),
    )
    if r.status_code != 200:
        return jsonify(r.json()), r.status_code
    commits = r.json()

    def fetch_detail(sha):
        try:
            d = requests.get("https://api.github.com/repos/" + repo + "/commits/" + sha, headers=gh_headers())
            return sha, (d.json() if d.status_code == 200 else None)
        except requests.RequestException:
            return sha, None

    details_by_sha = {}
    with ThreadPoolExecutor(max_workers=10) as ex:
        for sha, detail in ex.map(fetch_detail, [c["sha"] for c in commits]):
            details_by_sha[sha] = detail

    stats = {}
    timeline = []
    for c in commits:
        author = c.get("author")
        login = author["login"] if author else (c["commit"]["author"]["name"] if c["commit"].get("author") else "unknown")
        avatar = author["avatar_url"] if author else None
        date = c["commit"]["author"]["date"] if c["commit"].get("author") else None
        if login not in stats:
            stats[login] = {"login": login, "avatar_url": avatar, "commits": 0, "additions": 0, "deletions": 0, "test_commits": 0, "active_days": set()}
        stats[login]["commits"] += 1
        if date:
            stats[login]["active_days"].add(date[:10])

        d = details_by_sha.get(c["sha"])
        commit_additions = commit_deletions = 0
        if d:
            s = d.get("stats", {})
            commit_additions = s.get("additions", 0)
            commit_deletions = s.get("deletions", 0)
            stats[login]["additions"] += commit_additions
            stats[login]["deletions"] += commit_deletions
            touched_files = [f.get("filename", "") for f in d.get("files", [])]
            if any("test" in fn.lower() for fn in touched_files):
                stats[login]["test_commits"] += 1

        if date:
            timeline.append({"login": login, "date": date[:10], "additions": commit_additions, "deletions": commit_deletions})

    for v in stats.values():
        v["avg_lines_per_commit"] = round((v["additions"] + v["deletions"]) / v["commits"], 1) if v["commits"] else 0
        v["test_coverage_pct"] = round(100 * v["test_commits"] / v["commits"], 1) if v["commits"] else 0
        v["active_days_count"] = len(v["active_days"])
        del v["active_days"]

    return jsonify({"contributors": sorted(stats.values(), key=lambda x: -x["commits"]), "timeline": timeline})


TEAM_BUG_JOBS = {}


def run_team_bugs_job(job_id, repo, use_ai, branch=None):
    job = TEAM_BUG_JOBS[job_id]
    try:
        if not ensure_scanner_available():
            job["status"] = "error"
            job["error"] = "Could not find or clone the scanner tool"
            return

        target_dir = get_repo_dir(repo, branch)
        job["log"].append(f"Target ready: {target_dir}" + (f" (branch: {branch})" if branch else ""))

        import tempfile
        tmpdir = tempfile.mkdtemp()
        json_out = str(Path(tmpdir) / "report.json")
        cmd = [PYTHON_CMD, str(OTHER_SCANNER_PATH), str(target_dir), "--json", json_out, "--scan", "code", "--fail-on", "none"]
        if use_ai:
            cmd.append("--ai")

        scan_env = os.environ.copy()
        scan_env["PYTHONIOENCODING"] = "utf-8"
        proc = subprocess.Popen(
            cmd, cwd=OTHER_SCANNER_PATH.parent,
            stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
            text=True, bufsize=1, encoding="utf-8", errors="replace",
            env=scan_env,
        )
        for line in proc.stdout:
            line = line.rstrip()
            if line:
                job["log"].append(line)
        proc.wait()

        try:
            with open(json_out, "r", encoding="utf-8") as f:
                report = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            job["status"] = "error"
            job["error"] = "Scan did not produce a valid report."
            return

        job["log"].append("Attributing findings via git blame...")
        severity_rank = {"critical": 3, "high": 2, "medium": 1, "low": 0}
        contributors = {}
        for finding in report.get("findings", []):
            file_path = finding.get("file", "")
            fline = finding.get("line")
            if not file_path or not fline:
                continue
            if file_path.lower().endswith((".md", ".rst", ".txt")):
                continue
            if use_ai and finding.get("ai_verdict") == "false_positive":
                continue

            blame = subprocess.run(
                ["git", "blame", "-L", f"{fline},{fline}", "--porcelain", "--", file_path],
                cwd=target_dir, capture_output=True, text=True,
            )
            author = "unknown"
            for bline in blame.stdout.splitlines():
                if bline.startswith("author "):
                    author = bline[len("author "):].strip()
                    break

            if author not in contributors:
                contributors[author] = {"author": author, "total_bugs": 0, "by_severity": {"critical": 0, "high": 0, "medium": 0, "low": 0}, "worst": "low", "findings": []}
            c = contributors[author]
            c["total_bugs"] += 1
            sev = finding.get("severity", "low")
            if sev in c["by_severity"]:
                c["by_severity"][sev] += 1
            if severity_rank.get(sev, 0) > severity_rank.get(c["worst"], 0):
                c["worst"] = sev
            c["findings"].append({
                "file": file_path,
                "line": fline,
                "severity": sev,
                "description": finding.get("description", ""),
                "evidence": finding.get("evidence", ""),
                "impact": finding.get("impact", ""),
                "improvement": finding.get("improvement", ""),
            })

        job["status"] = "done"
        job["result"] = sorted(contributors.values(), key=lambda x: -x["total_bugs"])
    except Exception as e:
        job["status"] = "error"
        job["error"] = str(e)


@app.route("/api/team-summary", methods=["POST"])
def team_summary():
    """Generates a short, plain-English executive summary from
    aggregated team stats, using the same LLM setup as the rest of
    this tool. Purely descriptive; never used for any gating."""
    data = request.get_json(force=True)
    client, model = get_client()
    if not client:
        return jsonify({"error": "LLM_API_KEY not configured on the server."}), 400

    prompt = (
        "You are summarizing a software team's activity for a manager/executive. "
        "Write 2-3 short sentences, plain English, no markdown headers, no bullet points. "
        "Be factual and neutral, note both progress and any risk. Here is the data:\n\n"
        + json.dumps(data, indent=2)
    )
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You write brief, factual team-activity summaries for a manager. 2-3 sentences max, no formatting."},
                {"role": "user", "content": prompt},
            ],
        )
        return jsonify({"summary": response.choices[0].message.content})
    except Exception as e:
        return jsonify({"error": str(e)}), 502


@app.route("/api/team-report", methods=["POST"])
def team_report():
    """Builds a downloadable Team Performance report (PDF/Word/Markdown/
    HTML) from the Monitor tab's already-computed contributor data and
    executive summary - reuses the same builders as the vuln scan report."""
    data = request.get_json(force=True)
    fmt = (data.get("format") or "markdown").lower()
    contributors = data.get("contributors", [])
    summary = data.get("summary", "")
    range_label = data.get("range", "")

    lines = []
    if summary:
        lines += ["## Executive Summary", "", summary, ""]
    lines += ["## Contributor Breakdown", ""]
    for c in contributors:
        lines.append(f"### {c.get('login', 'unknown')}")
        lines.append("")
        lines.append(f"- **Commits:** {c.get('commits', 0)}")
        lines.append(f"- **Lines:** +{c.get('additions', 0)} / -{c.get('deletions', 0)}")
        lines.append(f"- **Test coverage:** {c.get('test_coverage_pct', 0)}%")
        lines.append(f"- **Active days:** {c.get('active_days_count', 0)}")
        lines.append(f"- **Bugs found:** {c.get('total_bugs', 0)} (worst severity: {c.get('worst', 'low')})")
        lines.append("")
    report_md = normalize_report_text("\n".join(lines))

    title = "Team Performance Report"
    subtitle = (range_label + " &middot; " if range_label else "") + \
        "generated " + datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    subtitle_plain = subtitle.replace("&middot;", "-")

    if fmt == "markdown":
        return Response(report_md, mimetype="text/markdown",
                         headers={"Content-Disposition": "attachment; filename=team-report.md"})
    if fmt == "html":
        return Response(build_report_html(report_md, title, subtitle), mimetype="text/html",
                         headers={"Content-Disposition": "attachment; filename=team-report.html"})
    if fmt == "docx":
        return Response(build_report_docx(report_md, title, subtitle_plain),
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         headers={"Content-Disposition": "attachment; filename=team-report.docx"})
    if fmt == "pdf":
        return Response(build_report_pdf(report_md, title, subtitle_plain), mimetype="application/pdf",
                         headers={"Content-Disposition": "attachment; filename=team-report.pdf"})
    return jsonify({"error": f"Unknown format: {fmt}"}), 400


@app.route("/api/team-bugs/start", methods=["POST"])
def team_bugs_start():
    """Starts the bug-attribution scan as a background job (mirrors
    /api/vuln-scan/start), so the Monitor tab can show live progress
    instead of blocking, especially when AI verification is enabled."""
    data = request.get_json(force=True)
    repo = data.get("repo", "").strip()
    use_ai = data.get("ai", False)
    branch = (data.get("branch") or "").strip() or None
    if not repo:
        return jsonify({"error": "repo required"}), 400

    job_id = str(uuid.uuid4())
    TEAM_BUG_JOBS[job_id] = {"status": "running", "log": [], "result": None, "error": None, "started": time.time()}

    thread = threading.Thread(target=run_team_bugs_job, args=(job_id, repo, use_ai, branch), daemon=True)
    thread.start()

    return jsonify({"job_id": job_id})


@app.route("/api/team-bugs/status/<job_id>")
def team_bugs_status(job_id):
    job = TEAM_BUG_JOBS.get(job_id)
    if not job:
        return jsonify({"error": "Unknown job_id"}), 404
    return jsonify({
        "status": job["status"],
        "log": job["log"],
        "result": job["result"],
        "error": job["error"],
        "elapsed": round(time.time() - job["started"], 1),
    })


@app.route("/api/branches")
def branches():
    """Lists ALL remote branches for the given repo (cloning/fetching it
    into a local cache first), so the dashboard can offer a GitHub-style
    base/compare selector for whichever repo is currently selected."""
    repo = request.args.get("repo", "").strip()
    cwd = get_repo_dir(repo)
    result = subprocess.run(
        ["git", "branch", "-r", "--format=%(refname:short)"], capture_output=True, text=True, cwd=cwd
    )
    current = subprocess.run(["git", "rev-parse", "--abbrev-ref", "HEAD"], capture_output=True, text=True, cwd=cwd)
    branch_list = [
        b.strip() for b in result.stdout.splitlines()
        # "origin/main" etc only - excludes the bare "origin" line git prints
        # for the remote's symbolic HEAD ref (origin/HEAD -> origin/main),
        # which isn't a real branch you can check out.
        if b.strip() and "/" in b.strip() and not b.strip().endswith("/HEAD")
    ]
    return jsonify({"branches": branch_list, "current": current.stdout.strip()})


@app.route("/api/current-diff")
def current_diff():
    """Runs git diff for the given repo (cloning/fetching it into a
    local cache first) and returns it, so the dashboard can get a diff
    for whichever repo is selected, not just the local AI-PR-Summary
    checkout. With no ?base=, returns uncommitted changes only (only
    meaningful for the AI-PR-Summary repo this server lives in).
    With ?base=main&compare=my-branch, returns the full diff between
    the two named branches, matching what a real PR would contain."""
    repo = request.args.get("repo", "").strip()
    base = request.args.get("base", "").strip()
    compare = request.args.get("compare", "HEAD").strip() or "HEAD"
    cwd = get_repo_dir(repo)
    if base:
        result = subprocess.run(
            ["git", "diff", f"{base}...{compare}"], capture_output=True, text=True, cwd=cwd
        )
    else:
        result = subprocess.run(["git", "diff"], capture_output=True, text=True, cwd=cwd)
    return jsonify({"diff": result.stdout})


@app.route("/api/scan", methods=["POST"])
def scan():
    """Runs the real summarize_pr.py pipeline: filter -> truncate -> LLM."""
    data = request.get_json(force=True)
    diff_text = data.get("diff", "")

    if not diff_text.strip():
        return jsonify({"error": "Diff is empty."}), 400

    filtered = filter_diff(diff_text)
    if not filtered.strip():
        return jsonify({"error": "Diff only contained ignored files (lockfiles, minified assets)."}), 400

    truncated, omitted = smart_truncate(filtered, MAX_CHARS)

    client, model = get_client()
    if not client:
        return jsonify({"error": "LLM_API_KEY not configured on the server."}), 400

    try:
        summary = call_llm(client, model, truncated)
    except Exception as e:
        return jsonify({"error": str(e)}), 502

    return jsonify({"markdown": summary, "omitted_files": omitted})


@app.route("/api/security-scan", methods=["POST"])
def security_scan():
    """Runs the deterministic security scanner on the full raw diff."""
    data = request.get_json(force=True)
    diff_text = data.get("diff", "")

    if not diff_text.strip():
        return jsonify({"error": "Diff is empty."}), 400

    findings, baselined = scan_security(diff_text)
    report = format_security_report(findings, baselined)

    return jsonify({"markdown": report, "findings": findings, "baselined": baselined})


def gh_headers():
    token = os.environ.get("GITHUB_TOKEN")
    return {"Authorization": "Bearer " + token} if token else {}


@app.route("/api/github/pulls")
def github_pulls():
    repo = request.args.get("repo", "")
    if "/" not in repo:
        return jsonify({"error": "repo must be owner/repo"}), 400
    r = requests.get(
        "https://api.github.com/repos/" + repo + "/pulls",
        params={"state": "all", "per_page": 10, "sort": "updated", "direction": "desc"},
        headers=gh_headers(),
    )
    return jsonify(r.json()), r.status_code


@app.route("/api/github/comments")
def github_comments():
    repo = request.args.get("repo", "")
    number = request.args.get("number", "")
    r = requests.get(
        "https://api.github.com/repos/" + repo + "/issues/" + number + "/comments",
        params={"per_page": 20},
        headers=gh_headers(),
    )
    return jsonify(r.json()), r.status_code


@app.route("/api/github/commits")
def github_commits():
    repo = request.args.get("repo", "")
    r = requests.get(
        "https://api.github.com/repos/" + repo + "/commits",
        params={"per_page": 10},
        headers=gh_headers(),
    )
    return jsonify(r.json()), r.status_code


@app.route("/api/github/commit/<sha>")
def github_commit_detail(sha):
    repo = request.args.get("repo", "")
    r = requests.get(
        "https://api.github.com/repos/" + repo + "/commits/" + sha,
        headers=gh_headers(),
    )
    return jsonify(r.json()), r.status_code


OTHER_SCANNER_REPO = "https://github.com/Alkhanjary/security-scan.git"
OTHER_SCANNER_PATH = Path(__file__).resolve().parent.parent / "security-scan" / "scanner.py"


def ensure_scanner_available():
    """If the security-scan tool isn't present locally, clone it fresh
    from GitHub. Keeps the two tools in separate folders (never merged),
    but self-heals if the local copy is missing (new machine, deleted
    folder, etc.) instead of just failing."""
    if OTHER_SCANNER_PATH.exists():
        return True
    scanner_dir = OTHER_SCANNER_PATH.parent
    if not scanner_dir.exists():
        subprocess.run(
            ["git", "clone", OTHER_SCANNER_REPO, str(scanner_dir)],
            capture_output=True, text=True,
        )
    return OTHER_SCANNER_PATH.exists()


import threading
import time
import uuid

VULN_JOBS = {}

# Each finished job keeps its whole log and result in memory. Without a cap a
# long-lived server accumulates every scan it ever ran (an AI scan of a big
# repo is tens of thousands of log lines plus a full findings list), which
# slowly starves the process. Completed scans are already persisted to
# scan_history.json, so trimming the in-memory copies loses nothing durable.
MAX_KEPT_JOBS = 20

# Hard ceiling on a single job's log. A misbehaving auto-started app can spew
# output indefinitely; without this, one runaway process could exhaust memory.
MAX_JOB_LOG_LINES = 20000


def _prune_vuln_jobs():
    """Drops the oldest finished jobs once we're over the cap. Running jobs are
    never pruned - something is still writing to them."""
    finished = [(j.get("started", 0), jid) for jid, j in VULN_JOBS.items()
                if j.get("status") != "running"]
    if len(VULN_JOBS) <= MAX_KEPT_JOBS:
        return
    finished.sort()
    for _, jid in finished[:len(VULN_JOBS) - MAX_KEPT_JOBS]:
        VULN_JOBS.pop(jid, None)


# Saved scan history lives OUTSIDE the repo, in the user's data directory.
#
# It used to sit next to server.py, which meant scanning this project scanned
# the tool's own saved findings: every stored `evidence` string is a verbatim
# line of risky-looking source from some earlier scan, so the scanner happily
# reported eval(), innerHTML and SQL concatenation "in scan_history.json". On
# the last run that produced six bogus findings and the AI narrative opened its
# remediation plan with "replace eval() at scan_history.json:89".
_LEGACY_SCAN_HISTORY_FILE = Path(__file__).resolve().parent / "scan_history.json"
SCAN_HISTORY_DIR = Path(
    os.environ.get("AI_PR_SUMMARY_DATA_DIR")
    or (Path(os.environ.get("LOCALAPPDATA") or Path.home() / ".local" / "share") / "ai-pr-summary")
)
SCAN_HISTORY_FILE = SCAN_HISTORY_DIR / "scan_history.json"
_scan_history_lock = threading.Lock()


def _migrate_legacy_scan_history():
    """Moves a pre-existing in-repo history file to the data directory once, so
    upgrading doesn't silently lose saved scans."""
    try:
        if not _LEGACY_SCAN_HISTORY_FILE.exists() or SCAN_HISTORY_FILE.exists():
            return
        SCAN_HISTORY_DIR.mkdir(parents=True, exist_ok=True)
        SCAN_HISTORY_FILE.write_bytes(_LEGACY_SCAN_HISTORY_FILE.read_bytes())
        _LEGACY_SCAN_HISTORY_FILE.unlink()
    except Exception:
        pass


def _load_scan_history():
    try:
        _migrate_legacy_scan_history()
        if SCAN_HISTORY_FILE.exists():
            return json.loads(SCAN_HISTORY_FILE.read_text(encoding="utf-8"))
    except Exception:
        pass
    return []


def _save_scan_history(records):
    try:
        SCAN_HISTORY_FILE.parent.mkdir(parents=True, exist_ok=True)
        SCAN_HISTORY_FILE.write_text(json.dumps(records, indent=2), encoding="utf-8")
    except Exception:
        pass


def _finding_key(f):
    """Stable identity for a finding across scans: file + category.
    We don't use line number because it shifts between commits."""
    return (
        (f.get("file") or "").strip(),
        (f.get("category") or f.get("description") or "")[:80].lower().strip(),
        (f.get("severity") or "").lower(),
    )


def save_and_compare_scan(repo_key, findings, scan_types, timestamp=None, files_scanned=None,
                          job_id=None):
    """Persists this scan to disk and compares it against the previous scan
    for the same repo_key. Returns a delta dict with new/fixed/recurring counts."""
    ts = timestamp or time.time()
    entry = {
        "repo": repo_key,
        "timestamp": ts,
        "name": "",
        # Recorded so a report can still be built from this scan after its
        # in-memory job is gone (server restart, or pruned past the job cap) -
        # the dashboard only knows the job_id it was handed at scan time.
        "job_id": job_id,
        "scan_types": list(scan_types or []),
        "files_scanned": files_scanned,
        # Keep the fields the report builder needs (evidence/impact/improvement),
        # not just the ones the trend comparison uses - otherwise a report
        # generated from history is missing the parts that make it useful.
        "findings": [
            {k: f.get(k) for k in (
                "file", "line", "severity", "category", "description", "scan_type",
                "evidence", "impact", "improvement", "source", "ai_verdict", "ai_reason",
            )}
            for f in (findings or [])
        ],
    }

    with _scan_history_lock:
        records = _load_scan_history()
        prev = next((r for r in reversed(records) if r.get("repo") == repo_key), None)
        records.append(entry)
        records = records[-200:]  # cap at 200 records total
        _save_scan_history(records)

    if not prev:
        return {"previous_scan": None, "new": len(findings or []), "fixed": 0, "recurring": 0,
                "trend": "first_scan", "previous_timestamp": None}

    prev_keys = {_finding_key(f) for f in (prev.get("findings") or [])}
    curr_keys = {_finding_key(f) for f in (findings or [])}
    new_count = len(curr_keys - prev_keys)
    fixed_count = len(prev_keys - curr_keys)
    recurring_count = len(curr_keys & prev_keys)

    if new_count == 0 and fixed_count == 0:
        trend = "unchanged"
    elif fixed_count > new_count:
        trend = "improving"
    elif new_count > fixed_count:
        trend = "worsening"
    else:
        trend = "mixed"

    return {
        "previous_scan": prev.get("timestamp"),
        "new": new_count,
        "fixed": fixed_count,
        "recurring": recurring_count,
        "trend": trend,
        "previous_total": len(prev.get("findings") or []),
        "current_total": len(findings or []),
    }


def find_free_port():
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.bind(("127.0.0.1", 0))
        return s.getsockname()[1]


def port_is_open(host, port, timeout=1.0):
    try:
        with socket.create_connection((host, port), timeout=timeout):
            return True
    except OSError:
        return False


def wait_for_any_port(host, candidate_ports, timeout_s, process):
    """Polls a list of candidate ports until one accepts a connection or the
    started process exits/times out. We can't know for certain which port an
    arbitrary app will bind to (many frameworks ignore $PORT), so we poll a
    short list of the most likely ones instead of guessing a single value."""
    deadline = time.time() + timeout_s
    while time.time() < deadline:
        if process.poll() is not None:
            return None
        for p in candidate_ports:
            if port_is_open(host, p):
                return p
        time.sleep(1)
    return None


def _read_text(path):
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except OSError:
        return ""


RUN_SCRIPT_EXTS = (".bat", ".cmd", ".sh", ".ps1")

# Words that mark a script as "this starts the app". Ranked - earlier is a
# stronger signal, so the most likely launcher gets tried first.
RUN_SCRIPT_HINTS = (
    "start", "run", "serve", "dev", "launch", "up", "boot", "app", "server", "web",
)

# Scripts that clearly do something other than start the app.
RUN_SCRIPT_EXCLUDE = (
    "test", "lint", "build-only", "deploy", "publish", "release", "install",
    "setup", "clean", "migrate", "seed", "backup", "stop", "kill", "uninstall",
)

SKIP_SCAN_DIRS = {
    "node_modules", ".git", "venv", ".venv", "env", "__pycache__", "dist",
    "build", ".next", ".nuxt", "target", "vendor", ".tox", "site-packages",
    "coverage", ".pytest_cache", ".mypy_cache", "bin", "obj",
}


def _find_run_scripts(target_dir, max_depth=6, limit=8, require_hint=True):
    """Recursively finds shell/batch scripts anywhere in the project, ranked
    most-likely-launcher first. Projects frequently ship a start.bat / run.sh
    that encodes the exact command needed, and in a monorepo it usually lives
    in a subfolder rather than the root.

    require_hint=True (the default, used to build start candidates) keeps only
    scripts whose name suggests they launch the app. require_hint=False returns
    EVERY script found - used for the AI snapshot, so the AI can read all of
    them and decide for itself which one actually starts the web app rather
    than trusting our name-based guess."""
    root = Path(target_dir)
    found = []
    for path in root.rglob("*"):
        if path.suffix.lower() not in RUN_SCRIPT_EXTS or not path.is_file():
            continue
        rel = path.relative_to(root)
        if len(rel.parts) > max_depth:
            continue
        if any(part in SKIP_SCAN_DIRS or part.startswith(".") for part in rel.parts[:-1]):
            continue
        stem = path.stem.lower()
        hit = next((i for i, h in enumerate(RUN_SCRIPT_HINTS) if h in stem), None)
        if require_hint:
            if any(bad in stem for bad in RUN_SCRIPT_EXCLUDE):
                continue
            if hit is None:
                continue
        # Rank: hint strength first (unhinted last), then shallower paths.
        rank = (hit if hit is not None else len(RUN_SCRIPT_HINTS), len(rel.parts), len(stem))
        found.append((rank, path))

    found.sort(key=lambda t: t[0])
    return [p for _, p in found[:limit]]


def detect_project(target_dir, skip_docker=False):
    """Looks at marker files in target_dir to produce an ordered list of
    candidate start configurations to try in sequence. Docker Compose comes
    first (most reliable for complex stacks), then language-specific
    heuristics, so the caller tries each and falls back on failure.
    Pass skip_docker=True to omit the Docker Compose entry."""
    target_dir = Path(target_dir)
    has = lambda name: (target_dir / name).exists()
    use_shell = os.name == "nt"
    candidates = []

    # --- Docker Compose (first choice when present) ---
    if not skip_docker:
        for compose_name in ("docker-compose.yml", "docker-compose.yaml", "compose.yml", "compose.yaml"):
            if has(compose_name):
                candidates.append({
                    "name": "Docker Compose",
                    "install": None,
                    "start": "docker compose up --build",
                    "env": {},
                    "candidate_ports": [3000, 5173, 8000, 8080, 4200, 80],
                    "use_shell": True,
                })
                break

    # --- Node.js / frontend (root or monorepo subfolder) ---
    _frontend_subdir = next(
        (target_dir / sub for sub in ("frontend", "client", "web", "ui", "app")
         if (target_dir / sub / "package.json").exists()),
        None,
    ) if not has("package.json") else None
    pkg_json_dir = _frontend_subdir if _frontend_subdir else target_dir
    if has("package.json") or _frontend_subdir:
        _fp = free_port = find_free_port()
        pkg = {}
        try:
            pkg = json.loads(_read_text(pkg_json_dir / "package.json") or "{}")
        except json.JSONDecodeError:
            pass
        scripts = pkg.get("scripts", {}) if isinstance(pkg, dict) else {}
        start_script = "dev" if "dev" in scripts else "start" if "start" in scripts else None
        _pkg_has = lambda name: (pkg_json_dir / name).exists()
        if _pkg_has("yarn.lock"):
            install, runner = ["yarn", "install"], ["yarn"]
        elif _pkg_has("pnpm-lock.yaml"):
            install, runner = ["pnpm", "install"], ["pnpm"]
        else:
            install, runner = ["npm", "install"], ["npm", "run"]
        if start_script:
            if runner[0] == "npm":
                start = ["npm", "run", start_script]
            else:
                start = runner + [start_script]
            candidates.append({
                "name": "Node.js" + (f" ({pkg_json_dir.name}/)" if _frontend_subdir else ""),
                "install": install,
                "start": start,
                "env": {"PORT": str(_fp)},
                "candidate_ports": [_fp, 3000, 5173, 8080, 4200, 8000],
                "use_shell": use_shell,
                "cwd": str(pkg_json_dir),
            })

    # --- Django ---
    if has("manage.py"):
        _fp = find_free_port()
        candidates.append({
            "name": "Django",
            "install": [PYTHON_CMD, "-m", "pip", "install", "-r", "requirements.txt"] if has("requirements.txt") else None,
            "start": [PYTHON_CMD, "manage.py", "runserver", f"127.0.0.1:{_fp}"],
            "env": {},
            "candidate_ports": [_fp],
            "use_shell": use_shell,
        })

    # --- Python (FastAPI / Flask / generic) ---
    py_entry_names = ("app.py", "main.py", "server.py", "run.py", "wsgi.py")
    if has("requirements.txt") or has("pyproject.toml") or any(has(c) for c in py_entry_names):
        _fp = find_free_port()
        req_text = (_read_text(target_dir / "requirements.txt") +
                    _read_text(target_dir / "pyproject.toml")).lower()
        entry = next((c for c in py_entry_names if has(c)), None)
        install = [PYTHON_CMD, "-m", "pip", "install", "-r", "requirements.txt"] if has("requirements.txt") else \
                  [PYTHON_CMD, "-m", "pip", "install", "-e", "."] if has("pyproject.toml") else None
        if "uvicorn" in req_text or "fastapi" in req_text:
            if entry:
                start = [PYTHON_CMD, "-m", "uvicorn", f"{Path(entry).stem}:app",
                         "--port", str(_fp), "--host", "127.0.0.1"]
            else:
                nested_pkg = next(
                    (p.parent.name for p in sorted(target_dir.rglob("main.py"))
                     if p.parent != target_dir and p.parent.parent == target_dir),
                    None,
                )
                module = f"{nested_pkg}.main" if nested_pkg else "main"
                start = [PYTHON_CMD, "-m", "uvicorn", f"{module}:app",
                         "--port", str(_fp), "--host", "127.0.0.1"]
        elif entry:
            start = [PYTHON_CMD, entry]
        else:
            start = None
        if start:
            candidates.append({
                "name": "Python (Flask/FastAPI/generic)",
                "install": install,
                "start": start,
                "env": {"PORT": str(_fp), "FLASK_RUN_PORT": str(_fp)},
                "candidate_ports": [_fp, 5000, 8000],
                "use_shell": use_shell,
            })

    # --- Ruby ---
    if has("Gemfile"):
        _fp = find_free_port()
        is_rails = has("config") and (target_dir / "config" / "application.rb").exists()
        candidates.append({
            "name": "Ruby" + (" on Rails" if is_rails else ""),
            "install": ["bundle", "install"],
            "start": ["bundle", "exec", "rails", "server", "-p", str(_fp)] if is_rails
                     else ["bundle", "exec", "rackup", "-p", str(_fp)],
            "env": {},
            "candidate_ports": [_fp, 3000],
            "use_shell": use_shell,
        })

    # --- PHP ---
    if has("composer.json") or has("index.php"):
        _fp = find_free_port()
        candidates.append({
            "name": "PHP",
            "install": ["composer", "install"] if has("composer.json") else None,
            "start": ["php", "-S", f"127.0.0.1:{_fp}"],
            "env": {},
            "candidate_ports": [_fp],
            "use_shell": use_shell,
        })

    # --- Go ---
    if has("go.mod"):
        _fp = find_free_port()
        candidates.append({
            "name": "Go",
            "install": ["go", "mod", "download"],
            "start": ["go", "run", "."],
            "env": {"PORT": str(_fp)},
            "candidate_ports": [_fp, 8080],
            "use_shell": use_shell,
        })

    # --- src-layout Python packages (e.g. AI-RMA/src/erma/main.py) ---
    # A package under a src/ directory can't be imported from the repo root at
    # all, so a correct-looking "uvicorn pkg.main:app" fails with
    # "No module named pkg". Running with cwd set to the package's parent (and
    # PYTHONPATH pointing there, which the caller adds) makes it importable
    # with no install step at all.
    for init in sorted(target_dir.rglob("__init__.py")):
        rel = init.relative_to(target_dir).parts
        if any(p in SKIP_SCAN_DIRS or p.endswith((".egg-info", ".dist-info")) for p in rel):
            continue
        if len(rel) > 5:
            continue
        pkg_dir = init.parent
        pkg_parent = pkg_dir.parent
        if pkg_parent == target_dir:
            continue  # already importable from the root; nothing to fix
        pkg_name = pkg_dir.name
        if not pkg_name.isidentifier():
            continue
        for entry in ("main.py", "app.py", "server.py", "asgi.py", "wsgi.py"):
            if not (pkg_dir / entry).exists():
                continue
            mod = f"{pkg_name}.{Path(entry).stem}"
            _fp = find_free_port()
            candidates.append({
                "name": f"Python package {mod} (src layout)",
                "cwd": str(pkg_parent),
                "install": None,
                "start": [PYTHON_CMD, "-m", "uvicorn", f"{mod}:app",
                          "--host", "127.0.0.1", "--port", str(_fp)],
                "env": {},
                "candidate_ports": [_fp, 8000],
                "use_shell": use_shell,
            })
            _fp2 = find_free_port()
            candidates.append({
                "name": f"Python entrypoint {pkg_name}/{entry}",
                "cwd": str(pkg_parent),
                "install": None,
                # Deliberately no 5000 here: that's this dashboard's own default
                # port, and guessing it invites detecting ourselves as "the app".
                "candidate_ports": [_fp2, 8000, 8080],
                "start": [PYTHON_CMD, str(pkg_dir / entry)],
                "env": {},
                "use_shell": use_shell,
            })
            break

    # --- Run scripts (.bat/.cmd/.sh/.ps1) found anywhere in the tree ---
    # Many projects ship a "start.bat"/"run.sh" that already encodes the exact
    # working directory, env vars and flags the app needs - often the only
    # thing that actually works. Search recursively, since in a monorepo the
    # launcher usually sits in a subfolder, not at the root.
    for script in _find_run_scripts(target_dir):
        candidates.append({
            "name": f"Run script ({script.name})",
            "cwd": str(script.parent),
            "install": None,
            "start": (str(script) if script.suffix.lower() in (".bat", ".cmd")
                      else f'powershell -ExecutionPolicy Bypass -File "{script}"'
                      if script.suffix.lower() == ".ps1"
                      else f'bash "{script}"'),
            "env": {},
            "candidate_ports": [3000, 5173, 8000, 8080, 4200, 5000, 80],
            "use_shell": True,
        })

    # --- Static HTML (last resort) ---
    static_dir = target_dir if has("index.html") else next(
        (target_dir / sub for sub in ("dashboard", "public", "dist", "build", "www", "static", "src")
         if (target_dir / sub / "index.html").exists()),
        None,
    )
    if static_dir:
        _fp = find_free_port()
        candidates.append({
            "name": "Static HTML",
            "install": None,
            "start": [PYTHON_CMD, "-m", "http.server", str(_fp), "--directory", str(static_dir)],
            "env": {},
            "candidate_ports": [_fp],
            "use_shell": use_shell,
        })

    return candidates or None


IGNORED_SCAN_DIRS = {".git", "node_modules", "__pycache__", ".venv", "venv", ".repo-cache",
                     "dist", "build", ".next", ".pytest_cache", "vendor", "target",
                     # Agent/editor config: not application code, and its
                     # allowlists are full of literal commands and localhost
                     # URLs that trip every pattern the scanner has.
                     ".claude", ".vscode", ".idea"}

# This tool's own output. A saved scan record stores each finding's `evidence`
# verbatim, so scanning a folder that contains one means re-detecting every
# risky-looking line the scanner ever saw, attributed to the JSON file.
SELF_ARTIFACT_FILES = {"scan_history.json"}


def _is_self_artifact(file_path):
    """True for a finding located in this tool's own output or in editor/agent
    config, neither of which is application code under review."""
    if not file_path:
        return False
    parts = str(file_path).replace("\\", "/").split("/")
    if parts[-1] in SELF_ARTIFACT_FILES:
        return True
    return any(p in IGNORED_SCAN_DIRS for p in parts[:-1])
MANIFEST_FILES = (
    "package.json", "requirements.txt", "pyproject.toml", "Pipfile", "Gemfile",
    "composer.json", "go.mod", "Cargo.toml", "pom.xml", "build.gradle",
    "Dockerfile", "docker-compose.yml", "docker-compose.yaml", "README.md",
)


def build_project_snapshot(target_dir, max_files=400, max_chars=18000):
    """Builds a compact description of the project - a capped file listing
    plus the content of any manifest/entrypoint-ish files present - for the
    AI project detector to reason about, without shipping the whole repo."""
    target_dir = Path(target_dir)
    paths = []
    for p in sorted(target_dir.rglob("*")):
        rel_parts = p.relative_to(target_dir).parts
        if any(part in IGNORED_SCAN_DIRS or part.endswith((".egg-info", ".dist-info")) for part in rel_parts):
            continue
        if p.is_file():
            paths.append(str(p.relative_to(target_dir)).replace("\\", "/"))
        if len(paths) >= max_files:
            break

    snippets = []
    remaining = max_chars
    seen_snippet_paths = set()

    def add_snippet(fp, head=None):
        """head caps how much of this one file is included, so a single huge
        entrypoint can't eat the whole budget that later files still need."""
        nonlocal remaining
        rel_path = str(fp.relative_to(target_dir)).replace("\\", "/")
        if rel_path in seen_snippet_paths or remaining <= 200:
            return
        seen_snippet_paths.add(rel_path)
        limit = min(remaining, head) if head else remaining
        text = _read_text(fp)[:limit]
        if not text.strip():
            return
        snippets.append(f"--- {rel_path} ---\n{text}")
        remaining -= len(text)

    # Root-level manifests first - the common case, and cheap to check.
    for name in MANIFEST_FILES:
        fp = target_dir / name
        if fp.exists() and fp.is_file():
            add_snippet(fp)

    # A monorepo commonly has these same manifest filenames nested
    # per-service instead of at the true root (e.g. backend/requirements.txt,
    # frontend/package.json) - the root-only lookup above would miss them
    # entirely, leaving the AI with no real dependency/framework evidence
    # even though it's right there one level down. Search the whole tree
    # for a few of each, capped so this can't blow the character budget.
    if remaining > 200:
        for name in MANIFEST_FILES:
            if remaining <= 200:
                break
            if name == "README.md":
                continue  # only the root README is worth the budget
            for fp in sorted(target_dir.rglob(name))[:2]:
                rel_parts = fp.relative_to(target_dir).parts
                if any(part in IGNORED_SCAN_DIRS for part in rel_parts) or not fp.is_file():
                    continue
                add_snippet(fp)
                if remaining <= 200:
                    break

    # Same idea for Dockerfiles/compose files, which don't have one fixed
    # name (Dockerfile.dev, docker-compose.prod.yml, etc.) so they need
    # their own glob patterns rather than an exact-name lookup.
    for pattern in ("Dockerfile*", "docker-compose*.yml", "docker-compose*.yaml", "compose*.yml", "compose*.yaml"):
        if remaining <= 200:
            break
        for fp in sorted(target_dir.rglob(pattern))[:3]:
            rel_parts = fp.relative_to(target_dir).parts
            if any(part in IGNORED_SCAN_DIRS for part in rel_parts) or not fp.is_file():
                continue
            add_snippet(fp)

    # EVERY script in the project (.bat/.cmd/.sh/.ps1), not just ones whose
    # name looks like a launcher. These usually spell out the exact command,
    # cwd and env the app needs, so their contents are the single most useful
    # evidence for "how do I run this" - and the AI reads all of them and
    # decides which actually starts the web app, rather than us pre-filtering
    # on filename and possibly hiding the one that matters.
    for fp in _find_run_scripts(target_dir, limit=40, require_hint=False):
        if remaining <= 200:
            break
        add_snippet(fp, head=1200)

    # Entrypoint source files - lets the AI see how the server is actually
    # created (app = FastAPI(), app.listen(PORT), if __name__ == "__main__",
    # hardcoded ports) rather than guessing from the filename alone.
    entry_names = ("main.py", "app.py", "server.py", "run.py", "wsgi.py", "asgi.py",
                   "manage.py", "index.js", "server.js", "app.js", "main.go", "main.rs")
    entry_hits = []
    for name in entry_names:
        for fp in sorted(target_dir.rglob(name))[:2]:
            rel_parts = fp.relative_to(target_dir).parts
            if any(part in IGNORED_SCAN_DIRS or part.endswith((".egg-info", ".dist-info"))
                   for part in rel_parts) or not fp.is_file():
                continue
            entry_hits.append(fp)
    # Shallowest first - the real entrypoint is rarely buried deep.
    entry_hits.sort(key=lambda p: len(p.relative_to(target_dir).parts))
    for fp in entry_hits[:6]:
        if remaining <= 200:
            break
        add_snippet(fp, head=2500)

    return "Files:\n" + "\n".join(paths) + "\n\n" + "\n\n".join(snippets)


def parse_json_response(raw):
    """Strips optional markdown code fences an LLM sometimes wraps JSON in,
    then parses it. Raises if the result still isn't valid JSON."""
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw, flags=re.MULTILINE)
    return json.loads(raw)


CODE_SCAN_EXTENSIONS = {
    ".py", ".js", ".jsx", ".ts", ".tsx", ".java", ".go", ".rb", ".php", ".rs",
    ".c", ".cpp", ".h", ".hpp", ".cs", ".html", ".yml", ".yaml", ".json", ".sh", ".ps1",
}


def list_source_files(target_dir, max_files=5000):
    """Collects EVERY matching source file (skipping noise dirs, binaries,
    absurdly large generated files) - no total-size cutoff, since the AI
    code scan batches these across as many calls as it takes instead of
    truncating the review."""
    target_dir = Path(target_dir)
    files = []
    for p in sorted(target_dir.rglob("*")):
        if any(part in IGNORED_SCAN_DIRS for part in p.relative_to(target_dir).parts):
            continue
        if not p.is_file() or p.suffix.lower() not in CODE_SCAN_EXTENSIONS:
            continue
        try:
            if p.stat().st_size > 300_000:
                continue
        except OSError:
            continue
        text = _read_text(p)
        if not text:
            continue
        files.append((str(p.relative_to(target_dir)).replace("\\", "/"), text))
        if len(files) >= max_files:
            break
    return files


def batch_files_by_char_budget(files, budget=30000):
    """Groups (path, text) pairs into batches that each fit a character
    budget, so a large codebase gets reviewed across multiple AI calls
    (each file staying whole - never split mid-file) instead of being cut
    off once one call's budget runs out."""
    batches = []
    current, current_size = [], 0
    for path, text in files:
        entry_size = len(text) + len(path) + 20
        if current and current_size + entry_size > budget:
            batches.append(current)
            current, current_size = [], 0
        current.append((path, text))
        current_size += entry_size
    if current:
        batches.append(current)
    return batches


def classify_scan_type(finding):
    """The regex scanner's own findings don't carry an explicit code/web/
    network tag, so infer one from the shape of the 'file' field it put the
    finding under: a URL means the web scan produced it, a bare host:port
    means the network scan did, anything else is a real source file (code)."""
    file_field = str(finding.get("file") or "")
    if file_field.startswith("http://") or file_field.startswith("https://"):
        return "web"
    if re.match(r"^[\w.\-]+:\d+$", file_field):
        return "network"
    return "code"


def ai_code_scan_run(target_dir, job):
    """Has the AI read EVERY matching source file and report vulnerabilities
    in its own judgement, instead of only matching the scanner's fixed regex
    rules. Batches the files across as many AI calls as it takes (each file
    kept whole, never split) so a large codebase gets fully reviewed instead
    of silently truncated to whatever fits one call's budget."""
    client, model = get_client()
    if not client:
        job["log"].append("AI code scan skipped: no LLM configured.")
        return [], 0
    files = list_source_files(target_dir)
    if not files:
        job["log"].append("AI code scan: no source files found to review.")
        return [], 0

    batches = batch_files_by_char_budget(files)
    job["log"].append(f"AI code scan: reviewing {len(files)} file(s) across {len(batches)} batch(es)...")

    all_findings = []
    for i, batch in enumerate(batches, 1):
        job["log"].append(f"AI code scan: batch {i}/{len(batches)} ({len(batch)} file(s): " +
                           ", ".join(p for p, _ in batch[:5]) + (", ..." if len(batch) > 5 else "") + ")")
        body = "\n\n".join(
            f"=== {path} ===\n" + "\n".join(f"{j + 1}: {line}" for j, line in enumerate(text.splitlines()))
            for path, text in batch
        )
        try:
            raw = call_llm(client, model, body, system_prompt=AI_CODE_SCAN_SYSTEM_PROMPT)
            findings = parse_json_response(raw)
            if isinstance(findings, list):
                all_findings += findings
        except Exception as e:
            job["log"].append(f"AI code scan batch {i}/{len(batches)} failed: {e}")

    for f in all_findings:
        f["source"] = "ai-code-scan"
        f["scan_type"] = "code"
        if isinstance(f.get("file"), str):
            f["file"] = re.sub(r"^===\s*|\s*===$", "", f["file"]).strip()
    job["log"].append(f"AI code scan found {len(all_findings)} finding(s) across {len(files)} file(s).")
    return all_findings, len(files)


# Hosts that must never be fetched, however the request arrives. Scanning
# private and loopback addresses is the tool's actual job, so blocking RFC1918
# wholesale would break it - but the cloud metadata services hand out IAM
# credentials to anything that can issue a plain GET, and no legitimate scan
# target lives there.
BLOCKED_PROBE_HOSTS = {
    "169.254.169.254",          # AWS / Azure / DigitalOcean IMDS
    "metadata.google.internal",  # GCP
    "metadata.goog",
    "100.100.100.100",          # Alibaba Cloud
    "fd00:ec2::254",            # AWS IMDSv6
}
MAX_PROBE_REDIRECTS = 5


def _probe_target_allowed(url):
    """Returns (ok, reason). Rejects non-HTTP schemes and cloud metadata hosts.

    Checked on every redirect hop as well as the initial URL: a server that
    answers with 302 -> http://169.254.169.254/ would otherwise walk straight
    past a check applied only to what the user typed."""
    try:
        parsed = urlparse(url)
    except ValueError:
        return False, "could not be parsed as a URL"
    if parsed.scheme.lower() not in ("http", "https"):
        return False, f"scheme '{parsed.scheme or 'none'}' is not allowed (http/https only)"
    host = (parsed.hostname or "").strip("[]").lower()
    if not host:
        return False, "has no host"
    if host in BLOCKED_PROBE_HOSTS:
        return False, "is a cloud metadata endpoint, which can hand out credentials"
    return True, ""


def probe_url(url, timeout=8):
    """Fetches a URL for the web scan.

    Redirects are followed manually so each hop can be re-validated; requests'
    own allow_redirects would follow a 302 into a blocked host without ever
    consulting the check above."""
    ok, why = _probe_target_allowed(url)
    if not ok:
        return {"url": url, "error": f"refused to fetch: URL {why}"}

    current = url
    try:
        for _ in range(MAX_PROBE_REDIRECTS + 1):
            r = requests.get(current, timeout=timeout, allow_redirects=False)
            if r.is_redirect or r.is_permanent_redirect:
                nxt = r.headers.get("Location")
                if not nxt:
                    break
                current = urljoin(current, nxt)
                ok, why = _probe_target_allowed(current)
                if not ok:
                    return {"url": url,
                            "error": f"refused to follow redirect to {current}: URL {why}"}
                continue
            return {"url": url, "final_url": current, "status": r.status_code,
                    "headers": dict(r.headers), "body_snippet": r.text[:1500]}
        return {"url": url, "error": f"stopped after {MAX_PROBE_REDIRECTS} redirects"}
    except requests.RequestException as e:
        return {"url": url, "error": str(e)}


def ai_web_scan_run(base_url, job):
    """AI-guided web scan: our code performs the actual HTTP requests, the
    AI decides what's worth checking next and interprets the results - it
    never touches the network itself."""
    client, model = get_client()
    if not client:
        job["log"].append("AI web scan skipped: no LLM configured.")
        return []
    common_paths = ["/", "/robots.txt", "/.well-known/security.txt", "/.env", "/.git/config", "/admin", "/api"]
    evidence = [probe_url(urljoin(base_url, p)) for p in common_paths]
    job["log"].append(f"AI web scan: probed {len(evidence)} path(s), asking AI what else to check...")

    extra_paths = []
    try:
        ask = call_llm(
            client, model,
            json.dumps(evidence, indent=2)[:20000] + "\n\nReturn additional_checks as instructed.",
            system_prompt=AI_WEB_SCAN_SYSTEM_PROMPT,
        )
        extra = parse_json_response(ask)
        if isinstance(extra, dict):
            extra_paths = (extra.get("additional_paths") or [])[:5]
    except Exception:
        pass

    if extra_paths:
        job["log"].append(f"AI requested {len(extra_paths)} more check(s): {extra_paths}")
        evidence += [probe_url(urljoin(base_url, p)) for p in extra_paths]

    try:
        raw = call_llm(client, model, json.dumps(evidence, indent=2)[:30000], system_prompt=AI_WEB_SCAN_SYSTEM_PROMPT)
        findings = parse_json_response(raw)
        if not isinstance(findings, list):
            findings = []
    except Exception as e:
        job["log"].append(f"AI web scan analysis failed: {e}")
        return []
    for f in findings:
        f["source"] = "ai-web-scan"
        f["scan_type"] = "web"
    job["log"].append(f"AI web scan found {len(findings)} finding(s).")
    return findings


COMMON_NETWORK_PORTS = [
    21, 22, 23, 25, 53, 80, 110, 111, 135, 139, 143, 443, 445, 993, 995,
    1433, 1521, 3000, 3306, 3389, 5432, 5900, 5984, 6379, 8000, 8080, 8443, 9200, 11211, 27017,
]


def simple_port_scan(host, ports=None, timeout=1.0, max_workers=30):
    ports = ports or COMMON_NETWORK_PORTS

    def check(port):
        try:
            with socket.create_connection((host, port), timeout=timeout) as s:
                banner = None
                try:
                    s.settimeout(1.0)
                    banner = s.recv(128).decode(errors="replace").strip() or None
                except OSError:
                    pass
                return {"port": port, "banner": banner}
        except OSError:
            return None

    open_ports = []
    with ThreadPoolExecutor(max_workers=max_workers) as ex:
        for res in ex.map(check, ports):
            if res:
                open_ports.append(res)
    return sorted(open_ports, key=lambda x: x["port"])


def ai_network_scan_run(host, job):
    """AI-guided network scan: our code performs the real port scan, the AI
    interprets which open ports/services are actually risky."""
    client, model = get_client()
    if not client:
        job["log"].append("AI network scan skipped: no LLM configured.")
        return []
    job["log"].append(f"AI network scan: probing common ports on {host}...")
    open_ports = simple_port_scan(host)
    job["log"].append(f"Found {len(open_ports)} open port(s).")
    if not open_ports:
        return []
    try:
        raw = call_llm(
            client, model, json.dumps({"host": host, "open_ports": open_ports}, indent=2),
            system_prompt=AI_NETWORK_SCAN_SYSTEM_PROMPT,
        )
        findings = parse_json_response(raw)
        if not isinstance(findings, list):
            findings = []
    except Exception as e:
        job["log"].append(f"AI network scan analysis failed: {e}")
        return []
    for f in findings:
        f["source"] = "ai-network-scan"
        f["scan_type"] = "network"
    job["log"].append(f"AI network scan found {len(findings)} finding(s).")
    return findings


# CLI tools that are normally installed as pip packages (not guaranteed to
# be on PATH as a standalone executable, even when the package itself is
# installed) - invoking them as "python -m <tool>" always works as long as
# the package is installed in that same Python environment, sidestepping
# PATH entirely. The AI detector doesn't know which "python" we'll actually
# run it with, so its raw suggestion gets normalized to PYTHON_CMD here.
PYTHON_MODULE_TOOLS = {
    "uvicorn", "gunicorn", "hypercorn", "flask", "pytest", "black", "mypy",
    "celery", "streamlit", "pip", "django-admin", "fastapi",
}


def normalize_python_command(cmd_str):
    if not isinstance(cmd_str, str) or not cmd_str.strip():
        return cmd_str
    parts = cmd_str.strip().split(None, 1)
    first = parts[0]
    rest = parts[1] if len(parts) > 1 else ""
    if first in PYTHON_MODULE_TOOLS:
        return f'"{PYTHON_CMD}" -m {first}' + (f" {rest}" if rest else "")
    if first in ("python", "python3", "py"):
        # Replace bare python/py/python3 with the actual interpreter running
        # this server, so the subprocess always finds a real Python executable
        # (avoids the Windows Store "python not found" alias and python3/py
        # availability differences across platforms).
        return f'"{PYTHON_CMD}"' + (f" {rest}" if rest else "")
    return cmd_str


def _resolve_candidate_cwd(target_dir, raw_cwd, job):
    """Maps the AI's repo-relative 'cwd' onto a real directory, refusing any
    path that escapes target_dir (the AI sees untrusted file names, so a
    '../..' can't be allowed to walk the host filesystem)."""
    root = Path(target_dir).resolve()
    if not raw_cwd or not isinstance(raw_cwd, str) or raw_cwd.strip() in (".", "./"):
        return root
    try:
        resolved = (root / raw_cwd.strip().replace("\\", "/")).resolve()
    except Exception:
        return root
    if root not in resolved.parents and resolved != root:
        job["log"].append(f"Ignoring out-of-tree cwd {raw_cwd!r}; using repo root.")
        return root
    if not resolved.is_dir():
        job["log"].append(f"AI cwd {raw_cwd!r} does not exist; using repo root.")
        return root
    return resolved


def _build_ai_candidates(raw_candidates, target_dir, job):
    """Turns the AI's raw candidate dicts into runnable detection dicts."""
    detections = []
    for c in raw_candidates:
        if not isinstance(c, dict) or not c.get("start_command"):
            continue
        run_cwd = _resolve_candidate_cwd(target_dir, c.get("cwd"), job)
        free_port = find_free_port()
        port_env_var = c.get("port_env_var")
        env = {port_env_var: str(free_port)} if port_env_var else {}
        likely_ports = [p for p in (c.get("likely_ports") or []) if isinstance(p, int)]
        detections.append({
            "name": c.get("framework") or "AI-detected project",
            "cwd": str(run_cwd),
            "install": normalize_python_command(c.get("install_command")),
            "start": normalize_python_command(c["start_command"]),
            "env": env,
            "candidate_ports": [free_port] + likely_ports,
            "use_shell": True,
        })
    return detections


# How many times the AI gets to look at the failures and correct itself.
# Each round costs one LLM call plus however long its commands take to fail,
# so this is capped - but 3 rounds is enough to work through the usual chain
# of "missing install" -> "wrong module path" -> "wrong port".
AI_RETRY_ROUNDS = 3


def ai_retry_project(target_dir, job, attempts, log_from):
    """Second chance for the AI: shows it the commands that were tried and the
    real error output they produced, and asks for corrected commands. This is
    what turns a one-shot guess into the AI actually diagnosing and fixing its
    own failure - the most common cause (wrong module path, missing install,
    tool not on PATH) is obvious from the error but invisible beforehand."""
    client, model = get_client()
    if not client or not attempts:
        return []

    report = []
    for a in attempts:
        report.append(f"--- Attempt: {a['name']} ---")
        report.append(f"cwd: {a['cwd']}")
        if a.get("install"):
            report.append(f"install command: {cmd_display(a['install'])}")
        report.append(f"start command: {cmd_display(a['start'])}")
        report.append(f"outcome: {a['outcome']}")
        if a.get("output"):
            report.append("output:\n" + "\n".join(a["output"][-40:]))
        report.append("")

    # Tell it plainly which runtimes exist on this machine, so it stops
    # proposing commands that can never run here and picks one that can.
    env_lines = []
    for tool in ("python", "py", "node", "npm", "yarn", "pnpm", "docker", "go", "ruby", "php", "java", "bash"):
        env_lines.append(f"  {tool}: {'available' if shutil.which(tool) else 'NOT INSTALLED'}")
    env_lines.append(f"  this server's Python interpreter: {PYTHON_CMD}")

    try:
        snapshot = build_project_snapshot(target_dir)
        payload = (
            snapshot
            + "\n\n=== TOOLS AVAILABLE ON THIS MACHINE ===\n" + "\n".join(env_lines)
            + "\n\n=== FAILED ATTEMPTS ===\n" + "\n".join(report)
        )[:MAX_CHARS]
        raw = call_llm(client, model, payload, system_prompt=PROJECT_RETRY_SYSTEM_PROMPT)
        info = parse_json_response(raw)
    except Exception as e:
        job["log"].append(f"AI retry analysis failed ({e}).")
        return []

    if info.get("diagnosis"):
        job["log"].append(f"AI diagnosis: {info['diagnosis']}")

    raw_candidates = info.get("candidates")
    if not isinstance(raw_candidates, list) or not raw_candidates:
        job["log"].append("AI could not suggest a corrected way to start this project.")
        return []

    detections = _build_ai_candidates(raw_candidates, target_dir, job)
    # Never re-run a command that already failed identically.
    already = {cmd_display(a["start"]).strip() for a in attempts}
    detections = [d for d in detections if cmd_display(d["start"]).strip() not in already]
    if not detections:
        job["log"].append("AI's corrected commands matched ones already tried — stopping.")
        return []

    job["log"].append(f"AI suggested {len(detections)} corrected command(s):")
    for i, d in enumerate(detections, 1):
        job["log"].append(f"  {i}. {d['name']} — {cmd_display(d['start'])}")
    return detections


def ai_detect_project(target_dir, job):
    """Uses the LLM to look at the actual project files (not a fixed list of
    marker files) and decide whether this is a runnable web app and every
    plausible way to install/start it. Returns a LIST of detection dicts
    (matching detect_project's shape), best-first, so the caller can try each
    until one actually serves a port. Empty list if it's not a web project or
    the call fails."""
    client, model = get_client()
    if not client:
        return []
    try:
        snapshot = build_project_snapshot(target_dir)
        raw = call_llm(client, model, snapshot, system_prompt=PROJECT_DETECT_SYSTEM_PROMPT)
        info = parse_json_response(raw)
    except Exception as e:
        job["log"].append(f"AI project detection failed ({e}); falling back to heuristics.")
        return []

    job["log"].append(f"AI project analysis: {info.get('reasoning', '')}")
    if not info.get("is_web_project"):
        job["log"].append("AI detection result: not a runnable web project — trying heuristics.")
        return []

    # Accept the new multi-candidate shape, but stay compatible with the old
    # flat single-command shape in case the model answers in that form.
    raw_candidates = info.get("candidates")
    if not isinstance(raw_candidates, list) or not raw_candidates:
        if info.get("start_command"):
            raw_candidates = [{
                "framework": info.get("framework"),
                "cwd": info.get("cwd"),
                "install_command": info.get("install_command"),
                "start_command": info.get("start_command"),
                "port_env_var": info.get("port_env_var"),
                "likely_ports": info.get("likely_ports"),
            }]
        else:
            job["log"].append("AI detected a web project but produced no start command — trying heuristics.")
            return []

    detections = _build_ai_candidates(raw_candidates, target_dir, job)

    if not detections:
        job["log"].append("AI returned no usable start commands — trying heuristics.")
        return []

    job["log"].append(f"AI proposed {len(detections)} way(s) to start this project:")
    for i, d in enumerate(detections, 1):
        where = Path(d["cwd"]).name if Path(d["cwd"]) != Path(target_dir).resolve() else "repo root"
        job["log"].append(f"  {i}. {d['name']} — {cmd_display(d['start'])}  (in {where})")
    return detections


def cmd_display(cmd):
    return cmd if isinstance(cmd, str) else " ".join(cmd)


# Env vars holding this tool's own credentials. The auto-build feature runs
# build/start commands that come from the scanned project itself (npm
# postinstall hooks, Dockerfiles, arbitrary start scripts), so that code is
# untrusted - handing it the server's full environment would leak the user's
# GitHub token and LLM API key to any repo they scan.
_SECRET_ENV_KEYS = {
    "LLM_API_KEY", "LLM_BASE_URL", "LLM_MODEL", "GITHUB_TOKEN",
    "OPENAI_API_KEY", "ANTHROPIC_API_KEY", "GH_TOKEN",
}


def _sanitized_child_env():
    """A copy of the environment with this tool's own secrets removed, for
    subprocesses whose command line originates from scanned/untrusted code."""
    env = os.environ.copy()
    for key in list(env):
        if key in _SECRET_ENV_KEYS or key.endswith(("_API_KEY", "_TOKEN", "_SECRET")):
            env.pop(key, None)
    return env


def stream_subprocess(cmd, cwd, env, shell, job, prefix="", timeout=None):
    """Runs cmd, appending every line of its output to job['log'] as it is
    produced (not just a summary once it finishes) - this is also what
    keeps the child's stdout pipe drained, which matters: a verbose build
    (npm install, docker build, etc.) can otherwise fill the OS pipe buffer
    and hang the whole subprocess, since nothing would be reading it.
    A background Timer enforces a hard wall-clock kill if given, since the
    read loop alone only checks between output lines and a stalled process
    with no output at all wouldn't otherwise be caught."""
    job["log"].append(f"$ {cmd_display(cmd)}")
    proc = subprocess.Popen(
        cmd, cwd=cwd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
        text=True, bufsize=1, env=env, shell=shell, errors="replace",
    )
    timer = None
    if timeout:
        timer = threading.Timer(timeout, proc.kill)
        timer.daemon = True
        timer.start()
    try:
        for line in proc.stdout:
            line = line.rstrip()
            if line:
                job["log"].append(prefix + line)
        proc.wait()
    finally:
        if timer:
            timer.cancel()
    return proc.returncode


def drain_process_stdout(process, job, prefix=""):
    """Keeps reading a long-running (Popen) process's stdout for its whole
    lifetime, appending each line to job['log'] live - used for the started
    app itself (e.g. 'docker compose up'), which keeps producing output
    long after the port opens and scanning begins."""
    try:
        for line in process.stdout:
            line = line.rstrip()
            if not line:
                continue
            # A started app can log forever (docker compose, a dev server with
            # file watching). Once past the ceiling, drop the oldest lines
            # instead of growing without bound - the tail is what matters.
            # log_offset counts dropped lines so the ?since= index the client
            # polls with stays absolute and never skips or repeats output.
            if len(job["log"]) >= MAX_JOB_LOG_LINES:
                drop = MAX_JOB_LOG_LINES // 4
                del job["log"][:drop]
                job["log_offset"] = job.get("log_offset", 0) + drop
                job["log"].append("  [log trimmed — keeping only the most recent output]")
            job["log"].append(prefix + line)
    except (ValueError, OSError):
        pass


def _base_executable(cmd):
    """Extracts the program name from a command (string or argv list), so we
    can check it exists before spending the full startup timeout waiting for
    a process that could never have run."""
    if isinstance(cmd, (list, tuple)):
        return str(cmd[0]) if cmd else ""
    if not isinstance(cmd, str) or not cmd.strip():
        return ""
    s = cmd.strip()
    if s.startswith('"'):
        end = s.find('"', 1)
        return s[1:end] if end > 0 else s[1:]
    return s.split()[0]


def _executable_available(cmd):
    """True if the command's program can actually be found/run."""
    exe = _base_executable(cmd)
    if not exe:
        return False
    # Shell builtins / compound commands - let the shell decide.
    if any(ch in exe for ch in "&|<>"):
        return True
    if os.path.sep in exe or (os.path.altsep and os.path.altsep in exe):
        return Path(exe).exists()
    return shutil.which(exe) is not None


MANIFEST_INSTALLS = [
    ("requirements.txt", lambda: [PYTHON_CMD, "-m", "pip", "install", "-r", "requirements.txt"]),
    ("pyproject.toml",   lambda: [PYTHON_CMD, "-m", "pip", "install", "-e", "."]),
    ("setup.py",         lambda: [PYTHON_CMD, "-m", "pip", "install", "-e", "."]),
]


def _find_python_manifest_dir(det, target_dir):
    """Finds the directory holding the dependency manifest for a Python start
    command: the entrypoint's nearest ancestor with a manifest, else the
    candidate's own cwd. Used to fill in an install step the AI left out."""
    root = Path(target_dir).resolve()
    run_cwd = Path(det.get("cwd") or target_dir).resolve()

    search_dirs = []
    # If the start command names a .py file, walk up from that file first -
    # in a monorepo the manifest sits beside the package, not at the repo root.
    parts = det.get("start") if isinstance(det.get("start"), (list, tuple)) else \
        shlex.split(det.get("start") or "", posix=False)
    for tok in parts:
        tok = tok.strip('"')
        if tok.lower().endswith(".py"):
            p = (run_cwd / tok)
            try:
                p = p.resolve()
            except Exception:
                continue
            d = p.parent
            while d == root or root in d.parents:
                search_dirs.append(d)
                if d == root:
                    break
                d = d.parent
            break
    search_dirs.append(run_cwd)

    for d in search_dirs:
        for fname, _ in MANIFEST_INSTALLS:
            if (d / fname).exists():
                return d, fname
    return None, None


def _ensure_install_step(det, target_dir, job):
    """The AI is told install_command must be non-null when a manifest exists,
    but it sometimes omits it anyway - which fails later as an unhelpful
    ModuleNotFoundError. If this is a Python candidate with no install step and
    we can find a real manifest, synthesize the install ourselves."""
    if det.get("install"):
        return det
    start_disp = cmd_display(det.get("start", ""))
    if "docker" in start_disp.lower():
        return det
    looks_python = (
        PYTHON_CMD in start_disp
        or re.search(r"\b(python3?|py)\b", start_disp)
        or ".py" in start_disp
        or "uvicorn" in start_disp or "gunicorn" in start_disp or "flask" in start_disp
    )
    if not looks_python:
        return det

    mdir, fname = _find_python_manifest_dir(det, target_dir)
    if not mdir:
        return det
    install = next(builder() for f, builder in MANIFEST_INSTALLS if f == fname)
    det = dict(det)
    det["install"] = install
    det["install_cwd"] = str(mdir)
    job["log"].append(
        f"No install step given for '{det['name']}' — added one from {fname} "
        f"in {mdir.name}/ so its dependencies are actually present."
    )
    return det


def _python_source_roots(target_dir, limit=12):
    """Finds directories that should be on PYTHONPATH: any directory holding an
    importable package (an immediate subdirectory with __init__.py).

    This is what makes a src-layout project work without installing it. For
    AI-RMA/src/erma/__init__.py the importable root is AI-RMA/src, so
    'import erma' only resolves if that directory is on the path - otherwise
    you get 'No module named erma' no matter how correct the command looks."""
    root = Path(target_dir)
    roots = []
    for init in root.rglob("__init__.py"):
        rel = init.relative_to(root).parts
        if any(part in SKIP_SCAN_DIRS or part.endswith((".egg-info", ".dist-info"))
               for part in rel):
            continue
        if len(rel) > 6:
            continue
        parent = init.parent.parent      # dir containing the package
        if parent not in roots:
            roots.append(parent)
    # Shallowest first, and always include the repo root itself as a fallback.
    roots.sort(key=lambda p: len(p.relative_to(root).parts))
    if root not in roots:
        roots.append(root)
    return roots[:limit]


# Flags that make a dev server spawn a watcher/child process. The parent stays
# alive even when the child dies on an import error, so a crash looks like a
# hang and we burn the whole startup timeout before learning anything. Stripping
# them makes the real error surface immediately, and auto-reload is useless here
# anyway since nothing edits the code mid-scan.
RELOAD_FLAGS = ("--reload", "--reload-dir", "--debug", "--dev")


def _expand_shell_vars(cmd, port):
    """Replaces POSIX-style ${VAR:-default} / $VAR / %VAR% placeholders with a
    concrete value. The AI likes emitting '--port ${PORT:-8000}', which Windows
    cmd passes through literally - uvicorn then dies on
    "'${PORT:-8000}' is not a valid integer". Substituting the port we actually
    chose makes the command work on any platform."""
    def sub_one(s):
        s = re.sub(r"\$\{PORT(?::-\d+)?\}", str(port), s, flags=re.IGNORECASE)
        s = re.sub(r"\$PORT\b", str(port), s, flags=re.IGNORECASE)
        s = re.sub(r"%PORT%", str(port), s, flags=re.IGNORECASE)
        # Any other ${VAR:-default} -> its default, which is the best guess
        # available and still beats passing the literal text through.
        s = re.sub(r"\$\{[A-Za-z_][A-Za-z0-9_]*:-([^}]*)\}", r"\1", s)
        return s

    if isinstance(cmd, (list, tuple)):
        return [sub_one(str(t)) for t in cmd]
    if isinstance(cmd, str):
        return sub_one(cmd)
    return cmd


def _strip_reload_flags(cmd):
    """Removes auto-reload/watcher flags from a start command."""
    if isinstance(cmd, (list, tuple)):
        out, skip_next = [], False
        for tok in cmd:
            if skip_next:
                skip_next = False
                continue
            if str(tok) == "--reload-dir":
                skip_next = True
                continue
            if str(tok) in RELOAD_FLAGS:
                continue
            out.append(tok)
        return type(cmd)(out) if not isinstance(cmd, tuple) else out
    if not isinstance(cmd, str):
        return cmd
    s = cmd
    s = re.sub(r"\s--reload-dir(?:=|\s+)\S+", "", s)
    for flag in RELOAD_FLAGS:
        s = re.sub(rf"\s{re.escape(flag)}(?=\s|$)", "", s)
    return s.strip()


# Errors that mean this attempt is already dead - no amount of extra waiting
# will produce a listening port, so bail out as soon as one appears.
FATAL_APP_PATTERNS = (
    r"ModuleNotFoundError",
    r"ImportError",
    r"SyntaxError",
    r"IndentationError",
    r"is not recognized as an internal or external command",
    r"command not found",
    r"Address already in use",
    r"only one usage of each socket address",
    r"Cannot find module",
    r"ERR_MODULE_NOT_FOUND",
    r"error TS\d+",
)
_FATAL_APP_RE = re.compile("|".join(FATAL_APP_PATTERNS), re.IGNORECASE)


def wait_for_port_or_fatal(host, candidate_ports, timeout_s, process, job, log_mark,
                           exclude_ports=()):
    """Like wait_for_any_port, but also gives up early when the app's own output
    shows an unrecoverable startup error. A watcher-based dev server keeps its
    parent process alive after the worker dies, so waiting on process exit alone
    means sitting out the full timeout on a failure we can already see.

    exclude_ports are ports that were ALREADY listening before we launched
    anything. They belong to some other process - frequently this dashboard's
    own server - and must never be reported as "the app came up", or we hand
    the scanner a completely unrelated target and every finding is about the
    wrong program."""
    exclude = set(exclude_ports)
    watch = [p for p in candidate_ports if p not in exclude]
    deadline = time.time() + timeout_s
    while time.time() < deadline:
        if process.poll() is not None:
            return None, "exited"
        for p in watch:
            if port_is_open(host, p):
                return p, None
        for line in job["log"][log_mark:]:
            if line.startswith("  [app] ") and _FATAL_APP_RE.search(line):
                return None, "fatal"
        time.sleep(1)
    return None, "timeout"


def _explain_startup_failure(job, log_mark):
    """Scans the failed attempt's own output for well-known startup errors and
    appends a plain-English cause + fix, so the log says what to do rather than
    leaving a raw traceback as the last word."""
    lines = job["log"][log_mark:]
    blob = "\n".join(lines)

    m = re.search(r"No module named ['\"]([\w.]+)['\"]", blob)
    if m:
        job["log"].append(f"Cause: the Python package '{m.group(1)}' isn't installed in this environment.")
        job["log"].append("  Fix: make sure the project's requirements.txt / pyproject.toml is installed "
                          "into the same Python that runs this server.")
        return
    if re.search(r"is not recognized as an internal or external command|command not found", blob):
        job["log"].append("Cause: the command's program isn't installed or isn't on PATH.")
        return
    m = re.search(r"ImportError: (.+)", blob)
    if m:
        job["log"].append(f"Cause: import failed — {m.group(1).strip()[:200]}")
        return
    if re.search(r"Address already in use|only one usage of each socket address", blob, re.I):
        job["log"].append("Cause: the port it wanted is already taken by another process.")
        return
    m = re.search(r"(SyntaxError|IndentationError): (.+)", blob)
    if m:
        job["log"].append(f"Cause: the app has a {m.group(1)} — {m.group(2).strip()[:160]}")
        return
    err_lines = [l for l in lines if re.search(r"error|exception|traceback", l, re.I)]
    if err_lines:
        job["log"].append(f"Last error seen: {err_lines[-1].strip()[:220]}")


def build_and_start_project(target_dir, job):
    """Installs deps and starts the detected project in the background.
    Tries AI-based detection first (reads the actual files to figure out
    what this project is, not a fixed list of frameworks), falling back to
    filename-heuristic detection if no LLM is configured or it fails.
    Returns (process, url) on success, or (None, None) if nothing could be
    detected/started (details are appended to job['log'] either way)."""
    ai_candidates = ai_detect_project(target_dir, job) or []
    heuristic_candidates = detect_project(target_dir) or []
    if not isinstance(heuristic_candidates, list):
        heuristic_candidates = [heuristic_candidates]

    # Build an ordered list: every AI-proposed command first (best-first), then
    # every heuristic candidate in priority order (Docker → Node → Python → …).
    # Each is tried until one actually serves a port.
    candidates = list(ai_candidates) + heuristic_candidates
    if not candidates:
        job["log"].append("Could not detect a runnable web project in this folder.")
        return None, None

    scripts = [c["name"] for c in heuristic_candidates if c["name"].startswith("Run script")]
    if scripts:
        job["log"].append(f"Found {len(scripts)} launcher script(s) in the project: "
                          + ", ".join(s[12:-1] for s in scripts))
    job["log"].append(f"Will try up to {len(candidates)} start method(s) until one works:")
    for i, c in enumerate(candidates, 1):
        job["log"].append(f"  {i}. {c['name']} — {cmd_display(c.get('start', ''))}")

    attempts = []   # what was tried + why it failed, fed back to the AI

    def _record(det, run_cwd, outcome, log_mark=None):
        attempts.append({
            "name": det.get("name", "?"),
            "cwd": str(run_cwd),
            "install": det.get("install"),
            "start": det.get("start"),
            "outcome": outcome,
            "output": [l for l in job["log"][log_mark:] if l.startswith("  [")] if log_mark else [],
        })

    py_source_roots = _python_source_roots(target_dir)

    def _try_start(det):
        det = _ensure_install_step(det, target_dir, job)
        run_cwd = Path(det.get("cwd") or target_dir)
        job["log"].append(f"Trying: {det['name']}" + (f" (from {run_cwd.name}/)" if run_cwd != Path(target_dir) else ""))
        run_env = _sanitized_child_env()
        run_env.update(det.get("env", {}))
        run_env["PYTHONIOENCODING"] = "utf-8"

        # Put every importable source root on PYTHONPATH. A src-layout project
        # (package under src/) can't be imported from the repo root at all
        # without this - the command looks right and still fails with
        # "No module named <pkg>". Costs nothing for non-Python candidates.
        if py_source_roots:
            existing = run_env.get("PYTHONPATH", "")
            joined = os.pathsep.join(str(p) for p in py_source_roots)
            run_env["PYTHONPATH"] = joined + (os.pathsep + existing if existing else "")

        # Auto-reload spawns a watcher whose parent survives the worker's death,
        # turning a visible crash into a silent 60s hang.
        stripped = _strip_reload_flags(det.get("start"))
        if cmd_display(stripped) != cmd_display(det.get("start")):
            job["log"].append("Removed auto-reload flag so startup errors surface immediately.")
            det = dict(det)
            det["start"] = stripped

        # Substitute ${PORT:-8000}-style placeholders, which Windows cmd would
        # otherwise pass through as literal text and the app would reject.
        want_port = (det.get("candidate_ports") or [find_free_port()])[0]
        expanded = _expand_shell_vars(det.get("start"), want_port)
        if cmd_display(expanded) != cmd_display(det.get("start")):
            job["log"].append(f"Substituted shell-style variables in the command (port {want_port}).")
            det = dict(det)
            det["start"] = expanded

        start_is_docker = "docker" in cmd_display(det.get("start", "")).lower()

        # Preflight: if the program itself isn't installed, nothing this
        # candidate does can work - skip it now instead of burning the whole
        # startup timeout waiting for a process that never ran.
        if not _executable_available(det.get("start")):
            exe = _base_executable(det.get("start"))
            job["log"].append(f"Skipping '{det['name']}': '{exe}' is not installed or not on PATH.")
            job["log"].append(f"  (install {exe}, or make sure it's on PATH for the account running this server)")
            _record(det, run_cwd, f"skipped: program '{exe}' is not installed or not on PATH")
            return None, None

        if det.get("install"):
            install_cwd = Path(det.get("install_cwd") or run_cwd)
            if not _executable_available(det["install"]):
                job["log"].append(
                    f"Skipping install: '{_base_executable(det['install'])}' is not on PATH — "
                    "trying the start command anyway.")
            else:
                job["log"].append("Installing dependencies..."
                                  + (f" (in {install_cwd.name}/)" if install_cwd != run_cwd else ""))
                try:
                    rc = stream_subprocess(det["install"], install_cwd, run_env, det["use_shell"], job,
                                           prefix="  [install] ", timeout=300)
                    if rc != 0:
                        job["log"].append(f"Install exited with code {rc} — trying the start command anyway "
                                          "(dependencies may already be present).")
                except (FileNotFoundError, OSError) as e:
                    job["log"].append(f"Could not install ({e}) — trying the start command anyway.")

        # Snapshot which candidate ports are ALREADY serving before we launch.
        # Anything listening now belongs to another process (very often this
        # dashboard itself on :5000), so treating it as "our app came up" would
        # point the scanner at the wrong program and make every finding wrong.
        preexisting = {p for p in det["candidate_ports"] if port_is_open("127.0.0.1", p)}
        if preexisting:
            job["log"].append(f"Ports already in use by something else, so they can't count "
                              f"as this app starting: {sorted(preexisting)}")
        usable = [p for p in det["candidate_ports"] if p not in preexisting]
        if not usable:
            job["log"].append(f"Skipping '{det['name']}': every candidate port is already "
                              "taken by another process, so a successful start can't be detected.")
            _record(det, run_cwd, "all candidate ports already occupied by other processes")
            return None, None

        job["log"].append("Starting: " + cmd_display(det["start"]))
        log_mark = len(job["log"])  # so we can inspect only this attempt's output
        try:
            process = subprocess.Popen(
                det["start"], cwd=run_cwd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                text=True, bufsize=1, env=run_env, shell=det["use_shell"], errors="replace",
            )
        except (FileNotFoundError, OSError) as e:
            job["log"].append(f"Could not start: {e}")
            _record(det, run_cwd, f"could not launch process: {e}", log_mark)
            return None, None

        threading.Thread(target=drain_process_stdout, args=(process, job, "  [app] "), daemon=True).start()

        startup_timeout = 240 if start_is_docker else 60
        job["log"].append(f"Waiting up to {startup_timeout}s for a port to open...")
        port, why = wait_for_port_or_fatal("127.0.0.1", det["candidate_ports"],
                                          timeout_s=startup_timeout, process=process,
                                          job=job, log_mark=log_mark,
                                          exclude_ports=preexisting)
        if not port:
            # Distinguish crashed / fatally-errored / alive-but-silent - they
            # point at completely different fixes.
            if why == "exited":
                job["log"].append(f"Process exited (code {process.returncode}) before serving a port.")
                outcome = f"process exited with code {process.returncode} before serving a port"
            elif why == "fatal":
                job["log"].append("Detected a fatal startup error in the app's output — "
                                  "stopping this attempt early instead of waiting for the timeout.")
                outcome = "fatal startup error in the app's own output (see output below)"
            else:
                job["log"].append(f"Timed out ({startup_timeout}s) — still running but no port opened. "
                                  f"Watched ports: {usable}")
                outcome = (f"still running after {startup_timeout}s but never opened any of "
                           f"ports {usable}")
                if preexisting:
                    # Very common when scanning an app that is already running:
                    # it wants a port another copy of itself is holding, so it
                    # can never bind. Say so, and point at the obvious fix.
                    job["log"].append(
                        f"Likely cause: this app wants port {sorted(preexisting)[0]}, which is already "
                        "taken by another process — probably the app itself, already running.")
                    job["log"].append(
                        f"  Fix: either stop whatever is on that port, or skip auto-build and put "
                        f"http://127.0.0.1:{sorted(preexisting)[0]} in the Target URL field directly.")
            _explain_startup_failure(job, log_mark)
            _record(det, run_cwd, outcome, log_mark)
            try:
                process.terminate()
            except Exception:
                pass
            return None, None

        url = f"http://127.0.0.1:{port}"
        job["log"].append(f"App is up at {url}")
        return process, url

    for i, det in enumerate(candidates):
        job["log"].append(f"--- Attempt {i + 1} of {len(candidates)}: {det['name']} ---")
        process, url = _try_start(det)
        if url:
            job["log"].append(f"Started successfully on attempt {i + 1} ({det['name']}).")
            return process, url
        if i < len(candidates) - 1:
            job["log"].append(f"Attempt {i + 1} failed — trying next option...")

    job["log"].append(f"All {len(candidates)} start method(s) failed.")

    # Everything we planned up front failed. Now let the AI keep working the
    # problem: each round it sees every command tried so far and the real error
    # output each produced, and proposes corrections. The most common causes
    # (wrong module path, missing install, tool not on PATH) are obvious in the
    # error output but were invisible when the first guesses were made, and a
    # fix often reveals the *next* problem - so we iterate rather than ask once.
    total_tried = len(candidates)
    for round_no in range(1, AI_RETRY_ROUNDS + 1):
        job["log"].append(f"=== AI retry round {round_no} of {AI_RETRY_ROUNDS}: "
                          "diagnosing the failures and correcting the commands ===")
        retry_candidates = ai_retry_project(target_dir, job, attempts, log_from=0)
        if not retry_candidates:
            break
        for i, det in enumerate(retry_candidates):
            total_tried += 1
            job["log"].append(f"--- Round {round_no}, retry {i + 1} of {len(retry_candidates)}: {det['name']} ---")
            process, url = _try_start(det)
            if url:
                job["log"].append(f"Started successfully on AI retry round {round_no} ({det['name']}).")
                return process, url
            job["log"].append(f"Retry {i + 1} failed — continuing...")

    job["log"].append(f"All {total_tried} start method(s) exhausted — could not bring the app up.")
    _summarize_start_failure(job, attempts)
    return None, None


def _summarize_start_failure(job, attempts):
    """Ends a failed auto-build with what the user should actually do, rather
    than leaving the last raw traceback as the final word in the log."""
    missing_tools, missing_modules = [], []
    for a in attempts:
        m = re.match(r"skipped: program '(.+?)' is not installed", a.get("outcome", ""))
        if m:
            missing_tools.append(m.group(1))
        for line in a.get("output", []):
            mm = re.search(r"No module named ['\"]([\w.]+)['\"]", line)
            if mm:
                missing_modules.append(mm.group(1))

    if missing_tools or missing_modules:
        job["log"].append("What to fix:")
    for tool in dict.fromkeys(missing_tools):
        hint = {"npm": " (install Node.js from https://nodejs.org)",
                "node": " (install Node.js from https://nodejs.org)",
                "docker": " (install Docker Desktop)",
                "yarn": " (install with: npm install -g yarn)",
                "pnpm": " (install with: npm install -g pnpm)"}.get(tool, "")
        job["log"].append(f"  - '{tool}' is not installed or not on PATH{hint}")
    for mod in dict.fromkeys(missing_modules):
        job["log"].append(f"  - Python package '{mod}' is missing — install the project's "
                          "requirements into the same Python that runs this server")
    job["log"].append("Tip: you can also start the app yourself and paste its URL into "
                      "the scan form instead of using auto-build.")


def stop_project(process, job):
    if not process:
        return
    try:
        process.terminate()
        process.wait(timeout=10)
    except subprocess.TimeoutExpired:
        process.kill()
    except Exception:
        pass
    job["log"].append("Stopped the auto-started app.")


def _warn_missing_manifests(target_dir, job, building):
    """A browser folder upload only carries files it could read as text, so
    manifests and compose files often never arrive - and then auto-build fails
    with a confusing 'no pyproject.toml found' about a file that plainly exists
    on disk. Say so up front rather than letting it look like a build bug."""
    tell = ("pyproject.toml", "requirements.txt", "setup.py", "package.json",
            "docker-compose.yml", "docker-compose.yaml", "compose.yml", "Dockerfile")
    found = {name for name in tell if any(target_dir.rglob(name))}
    if found:
        job["log"].append("Manifests present in upload: " + ", ".join(sorted(found)))
        return
    job["log"].append("NOTE: this upload contains no dependency manifest or compose file "
                      "(no pyproject.toml / requirements.txt / package.json / docker-compose.yml).")
    if building:
        job["log"].append("  Auto-build will likely fail: there is nothing to install from, and "
                          "Docker Compose can't be used without its compose file.")
        job["log"].append("  Browser folder uploads drop files they can't read as text. To keep the "
                          "project intact, use the 'Local folder path' source instead of Upload.")


def run_vuln_scan_job(job_id, repo, use_ai, scan_types, fail_on, include_test_files, files=None,
                       url=None, net_target=None, net_ports=None, auto_build=False,
                       local_path=None):
    job = VULN_JOBS[job_id]
    built_process = None
    try:
        import tempfile
        if local_path:
            target_dir = Path(local_path).expanduser().resolve()
            job["log"].append(f"Scanning folder in place (no upload): {target_dir}")
        elif files:
            upload_dir = Path(tempfile.mkdtemp())
            upload_root = upload_dir.resolve()
            rejected = 0
            for f in files:
                rel_path = f.get("path", "").lstrip("/\\")
                if not rel_path:
                    continue
                dest = upload_dir / rel_path
                # Path traversal guard: the client controls these paths, and
                # neither lstrip() nor pathlib stops "../.." from climbing out
                # of the temp dir (on Windows an absolute "C:/..." entry even
                # replaces the base entirely). Resolve and confirm the result
                # is still inside the upload dir before writing anything.
                try:
                    resolved = dest.resolve()
                    resolved.relative_to(upload_root)
                except (ValueError, OSError):
                    rejected += 1
                    continue
                resolved.parent.mkdir(parents=True, exist_ok=True)
                resolved.write_text(f.get("content", ""), encoding="utf-8", errors="replace")
            target_dir = upload_dir
            job["log"].append(f"Uploaded {len(files) - rejected} file(s) to scan: {target_dir}")
            if rejected:
                job["log"].append(f"Ignored {rejected} upload path(s) that pointed outside the scan folder.")
            _warn_missing_manifests(target_dir, job, auto_build and "web" in scan_types)
        else:
            job["log"].append(f"Cloning {repo}...")
            target_dir = get_repo_dir(repo)
            n_files = sum(1 for _ in target_dir.rglob("*") if _.is_file())
            job["log"].append(f"Target ready: {target_dir} ({n_files} file(s))")

        if "web" in scan_types and auto_build and not url:
            built_process, built_url = build_and_start_project(target_dir, job)
            if built_url:
                url = built_url

        # Stop clearly and say exactly what's missing, rather than letting
        # the external scanner abort the whole run with a generic message
        # that doesn't say which requested scan type actually caused it.
        missing_targets = []
        if "web" in scan_types and not url:
            missing_targets.append("web scan: no target URL (auto-build didn't produce one, and none was entered)")
        if "network" in scan_types and not net_target and not url:
            missing_targets.append("network scan: no host/IP available")
        if missing_targets:
            job["status"] = "error"
            job["error"] = "Scan stopped before running - could not proceed with: " + "; ".join(missing_targets)
            return

        tmpdir = tempfile.mkdtemp()
        json_out = str(Path(tmpdir) / "report.json")
        scan_arg = ",".join(scan_types) if scan_types else "code"
        cmd = [
            PYTHON_CMD, str(OTHER_SCANNER_PATH), str(target_dir),
            "--json", json_out,
            "--scan", scan_arg,
            "--fail-on", fail_on,
        ]
        if use_ai:
            cmd.append("--ai")
        if include_test_files:
            cmd.append("--include-test-files")
        if url:
            cmd += ["--url", url]
        if net_target:
            cmd += ["--net-target", net_target]
        if net_ports:
            cmd += ["--net-ports", net_ports]

        job["log"].append("Starting scan: " + " ".join(cmd))

        scan_env = os.environ.copy()
        scan_env["PYTHONIOENCODING"] = "utf-8"
        proc = subprocess.Popen(
            cmd, cwd=OTHER_SCANNER_PATH.parent,
            stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
            text=True, bufsize=1, encoding="utf-8", errors="replace",
            env=scan_env,
        )
        for line in proc.stdout:
            line = line.rstrip()
            if line:
                job["log"].append(line)
        proc.wait()

        try:
            with open(json_out, "r", encoding="utf-8") as f:
                report = json.load(f)

            # Drop detections inside this tool's own artifacts and editor/agent
            # config. The external scanner has no reason to know about those,
            # but they are pure noise: scan_history.json stores past findings'
            # evidence verbatim, so scanning it re-reports every risky-looking
            # line the tool has ever seen.
            raw_findings = report.get("findings") or []
            kept = [f for f in raw_findings if not _is_self_artifact(f.get("file"))]
            if len(kept) != len(raw_findings):
                job["log"].append(
                    f"Ignored {len(raw_findings) - len(kept)} detection(s) inside this tool's own "
                    "output or editor config (not application code).")
            report["findings"] = kept

            for finding in report["findings"]:
                finding.setdefault("scan_type", classify_scan_type(finding))

            if use_ai:
                ai_findings = []
                if "code" in scan_types:
                    code_findings, code_files = ai_code_scan_run(target_dir, job)
                    ai_findings += code_findings
                    report["files_scanned"] = max(report.get("files_scanned", 0), code_files)
                if "web" in scan_types and url:
                    ai_findings += ai_web_scan_run(url, job)
                if "network" in scan_types:
                    net_host = net_target.split("/")[0] if net_target else (urlparse(url).hostname if url else None)
                    if net_host:
                        ai_findings += ai_network_scan_run(net_host, job)
                if ai_findings:
                    job["log"].append(f"AI scan pass added {len(ai_findings)} additional finding(s) total.")
                    report["findings"] = (report.get("findings") or []) + ai_findings

                    # The regex scanner already set exit_code from its own findings;
                    # make sure AI-only findings can also trigger the gate instead of
                    # silently passing just because the regex pass alone found nothing.
                    severity_rank = {"low": 0, "medium": 1, "high": 2, "critical": 3}
                    if fail_on != "none" and not report.get("exit_code"):
                        threshold_rank = severity_rank.get(fail_on, 2)
                        gate_triggered = any(
                            severity_rank.get((f.get("severity") or "").lower(), -1) >= threshold_rank
                            and (include_test_files or not f.get("likely_test_fixture"))
                            for f in ai_findings
                        )
                        if gate_triggered:
                            report["exit_code"] = 2
                            job["log"].append(f"AI-found finding(s) alone triggered the gate (threshold: {fail_on}).")

            all_findings = report.get("findings") or []
            repo_key = repo or (local_path if local_path else "(uploaded)")
            delta = save_and_compare_scan(repo_key, all_findings, scan_types,
                                          files_scanned=report.get("files_scanned"),
                                          timestamp=job.get("started"), job_id=job_id)
            report["history_delta"] = delta
            job["status"] = "done"
            job["result"] = report
        except (FileNotFoundError, json.JSONDecodeError):
            job["status"] = "error"
            job["error"] = "Scanner did not produce a valid report."
    except Exception as e:
        job["status"] = "error"
        job["error"] = str(e)
    finally:
        stop_project(built_process, job)


@app.route("/api/vuln-scan/start", methods=["POST"])
@rate_limit("vuln_scan_start")
@require_csrf_token
def vuln_scan_start():
    """Starts the vulnerability scan as a background job and returns a
    job_id immediately, so the dashboard can poll for live progress
    instead of blocking on one long request (AI verification especially
    can take a while)."""
    data = request.get_json(force=True)
    repo = data.get("repo", "").strip()
    files = data.get("files")
    local_path = (data.get("local_path") or "").strip()
    if not ensure_scanner_available():
        return jsonify({"error": f"Could not find or clone the scanner tool (expected at {OTHER_SCANNER_PATH})"}), 500

    # Scanning a folder in place, instead of uploading it. A browser folder
    # upload only carries files it can read as text, so manifests and compose
    # files that auto-build depends on can silently go missing - reading the
    # folder directly keeps the project intact. Safe here because this server
    # is a localhost dev tool running as the same user who picks the path.
    if local_path:
        p = Path(local_path).expanduser()
        try:
            p = p.resolve(strict=True)
        except (OSError, RuntimeError):
            return jsonify({"error": f"That folder does not exist: {local_path}"}), 400
        if not p.is_dir():
            return jsonify({"error": f"Not a folder: {p}"}), 400

    job_id = str(uuid.uuid4())
    _prune_vuln_jobs()
    VULN_JOBS[job_id] = {
        "status": "running", "log": [], "result": None, "error": None, "started": time.time(),
        "repo": repo or (str(local_path) if local_path else "(uploaded folder)"), "name": "",
    }
    thread = threading.Thread(
        target=run_vuln_scan_job,
        args=(
            job_id, repo,
            data.get("ai", True),
            data.get("scan_types", ["code"]),
            data.get("fail_on", "high"),
            data.get("include_test_files", False),
            files,
        ),
        kwargs={
            "url": (data.get("url") or "").strip() or None,
            "net_target": (data.get("net_target") or "").strip() or None,
            "net_ports": (data.get("net_ports") or "").strip() or None,
            "auto_build": bool(data.get("auto_build", False)),
            "local_path": local_path or None,
        },
        daemon=True,
    )
    thread.start()

    return jsonify({"job_id": job_id})


@app.route("/api/vuln-scan/status/<job_id>")
def vuln_scan_status(job_id):
    """Poll a running scan. Pass ?since=N to get only log lines after index N -
    an AI scan of a large repo produces thousands of lines, and resending the
    whole log every second (plus re-rendering it client-side) is what made the
    dashboard crawl. The client keeps its own copy and appends the delta."""
    job = VULN_JOBS.get(job_id)
    if not job:
        return jsonify({"error": "Unknown job_id - the server may have restarted."}), 404

    log = job["log"]
    offset = job.get("log_offset", 0)   # lines already dropped off the front
    try:
        since = max(0, int(request.args.get("since", 0)))
    except (TypeError, ValueError):
        since = 0

    # `since` is an absolute line index. Translate it into a position in the
    # list we still hold; if the client is asking for lines already trimmed
    # away, give it everything we have left rather than a gap.
    start = min(max(since - offset, 0), len(log))
    delta = log[start:]

    return jsonify({
        "status": job["status"],
        "log": delta,
        "log_total": offset + len(log),   # absolute count, for the next ?since=
        "result": job["result"],
        "error": job["error"],
        "elapsed": round(time.time() - job["started"], 1),
    })


@app.route("/api/vuln-scan/history")
def vuln_scan_history():
    """Lists recent scans: in-memory session jobs first, then persistent
    history from disk (survives server restarts). Includes trend data."""
    # In-memory session jobs
    items = []
    for job_id, job in VULN_JOBS.items():
        entry = {
            "job_id": job_id, "repo": job.get("repo", ""), "status": job["status"],
            "started": job["started"], "source": "session", "name": job.get("name", ""),
            "scan_ref": f"job:{job_id}",
        }
        if job.get("result"):
            findings = job["result"].get("findings", [])
            counts = {"critical": 0, "high": 0, "medium": 0, "low": 0}
            for f in findings:
                sev = (f.get("severity") or "").lower()
                if sev in counts:
                    counts[sev] += 1
            entry["files_scanned"] = job["result"].get("files_scanned")
            entry["total_findings"] = len(findings)
            entry["counts"] = counts
            entry["history_delta"] = job["result"].get("history_delta")
        items.append(entry)
    items.sort(key=lambda e: -e["started"])

    # Persistent history (disk) — for scans from previous sessions
    with _scan_history_lock:
        disk_records = _load_scan_history()
    session_ts = {job["started"] for job in VULN_JOBS.values()}
    for rec in reversed(disk_records):
        if rec.get("timestamp") in session_ts:
            continue  # already shown from in-memory
        counts = {"critical": 0, "high": 0, "medium": 0, "low": 0}
        for f in (rec.get("findings") or []):
            sev = (f.get("severity") or "").lower()
            if sev in counts:
                counts[sev] += 1
        items.append({
            "repo": rec.get("repo", ""), "status": "done",
            "started": rec.get("timestamp", 0),
            "total_findings": len(rec.get("findings") or []),
            "counts": counts,
            "scan_types": rec.get("scan_types"),
            "source": "history",
            "name": rec.get("name", ""),
            "scan_ref": f"ts:{rec.get('timestamp', 0)}",
        })

    items.sort(key=lambda e: -e["started"])
    return jsonify({"items": items[:25]})


@app.route("/api/browse")
def browse_folders():
    """Lists sub-folders of a directory so the dashboard can offer a real
    folder picker for the 'scan a local folder' option.

    A browser deliberately never reveals the absolute path of a folder the user
    picks (webkitdirectory only yields relative paths), so the path has to be
    resolved server-side. Only directory NAMES are returned - never file
    contents - and the server already binds to loopback with a localhost-only
    CORS policy, so this doesn't widen what a local user can reach.
    """
    raw = (request.args.get("path") or "").strip()

    # No path yet: offer sensible starting points rather than a bare filesystem
    # root the user then has to click through.
    if not raw:
        home = Path.home()
        shortcuts = []
        for label, p in (("Home", home), ("Desktop", home / "Desktop"),
                         ("Documents", home / "Documents"), ("Downloads", home / "Downloads")):
            if p.is_dir():
                shortcuts.append({"name": label, "path": str(p)})
        if os.name == "nt":
            for letter in "CDEFGH":
                drive = Path(f"{letter}:\\")
                if drive.exists():
                    shortcuts.append({"name": f"{letter}: drive", "path": str(drive)})
        else:
            shortcuts.append({"name": "/", "path": "/"})
        return jsonify({"path": None, "parent": None, "shortcuts": shortcuts, "entries": []})

    try:
        current = Path(raw).expanduser().resolve(strict=True)
    except (OSError, RuntimeError):
        return jsonify({"error": f"No such folder: {raw}"}), 400
    if not current.is_dir():
        return jsonify({"error": f"Not a folder: {current}"}), 400

    entries, unreadable = [], False
    try:
        for child in sorted(current.iterdir(), key=lambda p: p.name.lower()):
            try:
                if not child.is_dir():
                    continue
            except OSError:
                continue  # broken junction / permission denied on stat
            if child.name.startswith(("$", ".")) and child.name not in (".github",):
                continue
            entries.append({"name": child.name, "path": str(child)})
    except PermissionError:
        unreadable = True

    parent = str(current.parent) if current.parent != current else None
    # Flag folders that look like a project, so the right one is easy to spot.
    markers = ("package.json", "pyproject.toml", "requirements.txt", "docker-compose.yml",
               "setup.py", "go.mod", "Cargo.toml", ".git")
    for e in entries:
        try:
            e["is_project"] = any((Path(e["path"]) / m).exists() for m in markers)
        except OSError:
            e["is_project"] = False

    return jsonify({
        "path": str(current),
        "parent": parent,
        "entries": entries,
        "unreadable": unreadable,
        "is_project": any((current / m).exists() for m in markers),
    })


@app.route("/api/vuln-scan/rename", methods=["PATCH"])
@require_csrf_token
def vuln_scan_rename():
    """Rename a scan. Body: {scan_ref: 'job:<id>' or 'ts:<ts>', name: '...'}"""
    data = request.get_json(force=True)
    ref  = (data.get("scan_ref") or "").strip()
    name = (data.get("name") or "").strip()[:120]
    if not ref:
        return jsonify({"error": "scan_ref required"}), 400

    if ref.startswith("job:"):
        job_id = ref[4:]
        if job_id not in VULN_JOBS:
            return jsonify({"error": "job not found"}), 404
        VULN_JOBS[job_id]["name"] = name
        return jsonify({"ok": True})

    if ref.startswith("ts:"):
        ts = float(ref[3:])
        with _scan_history_lock:
            records = _load_scan_history()
            updated = False
            for rec in records:
                if abs(rec.get("timestamp", 0) - ts) < 1:
                    rec["name"] = name
                    updated = True
                    break
            if not updated:
                return jsonify({"error": "scan not found in history"}), 404
            _save_scan_history(records)
        return jsonify({"ok": True})

    return jsonify({"error": "invalid scan_ref format"}), 400


@app.route("/api/vuln-scan/delete", methods=["DELETE"])
@require_csrf_token
def vuln_scan_delete():
    """Deletes a scan. Body: {scan_ref: 'job:<id>' or 'ts:<ts>'}.
    A session job is dropped from memory; a saved scan is removed from
    scan_history.json. Deleting a saved scan also drops its in-memory twin
    (same timestamp) so it can't reappear in the list."""
    data = request.get_json(force=True)
    ref = (data.get("scan_ref") or "").strip()
    if not ref:
        return jsonify({"error": "scan_ref required"}), 400

    if ref.startswith("job:"):
        job_id = ref[4:]
        job = VULN_JOBS.get(job_id)
        if not job:
            return jsonify({"error": "job not found"}), 404
        if job.get("status") == "running":
            return jsonify({"error": "That scan is still running — wait for it to finish first."}), 409
        started = job.get("started")
        VULN_JOBS.pop(job_id, None)
        # Drop the matching persisted record too, so it doesn't come straight
        # back on the next refresh as a "saved" entry.
        with _scan_history_lock:
            records = _load_scan_history()
            kept = [r for r in records if abs(r.get("timestamp", 0) - (started or -1)) >= 1]
            if len(kept) != len(records):
                _save_scan_history(kept)
        return jsonify({"ok": True})

    if ref.startswith("ts:"):
        try:
            ts = float(ref[3:])
        except ValueError:
            return jsonify({"error": "invalid timestamp in scan_ref"}), 400
        with _scan_history_lock:
            records = _load_scan_history()
            kept = [r for r in records if abs(r.get("timestamp", 0) - ts) >= 1]
            if len(kept) == len(records):
                return jsonify({"error": "scan not found in history"}), 404
            _save_scan_history(kept)
        for jid, j in list(VULN_JOBS.items()):
            if abs(j.get("started", 0) - ts) < 1:
                VULN_JOBS.pop(jid, None)
        return jsonify({"ok": True})

    return jsonify({"error": "invalid scan_ref format"}), 400


def _resolve_scan(ref):
    """Return {findings, name, repo, timestamp, counts} for a scan_ref."""
    if (ref or "").startswith("job:"):
        job_id = ref[4:]
        job = VULN_JOBS.get(job_id)
        if job and job.get("result"):
            findings = job["result"].get("findings", [])
            return {
                "findings": findings, "name": job.get("name", ""),
                "repo": job.get("repo", ""), "timestamp": job["started"],
                "counts": count_findings_by_severity(findings),
                "files_scanned": job["result"].get("files_scanned"),
            }
        # The job is gone (server restarted, or pruned past the cap) but the
        # scan itself was persisted - look it up by the job_id recorded there,
        # since the dashboard only ever knows the job_id it was given.
        with _scan_history_lock:
            records = _load_scan_history()
        rec = next((r for r in reversed(records) if r.get("job_id") == job_id), None)
        if not rec:
            return None
        findings = rec.get("findings", [])
        return {
            "findings": findings, "name": rec.get("name", ""),
            "repo": rec.get("repo", ""), "timestamp": rec.get("timestamp", 0),
            "counts": count_findings_by_severity(findings),
            "files_scanned": rec.get("files_scanned"),
        }
    if (ref or "").startswith("ts:"):
        ts = float(ref[3:])
        with _scan_history_lock:
            records = _load_scan_history()
        rec = next((r for r in records if abs(r.get("timestamp", 0) - ts) < 1), None)
        if not rec:
            return None
        findings = rec.get("findings", [])
        return {
            "findings": findings, "name": rec.get("name", ""),
            "repo": rec.get("repo", ""), "timestamp": rec.get("timestamp", 0),
            "counts": count_findings_by_severity(findings),
            "files_scanned": rec.get("files_scanned"),
        }
    return None


@app.route("/api/vuln-scan/compare", methods=["POST"])
@rate_limit("vuln_scan_compare")
@require_csrf_token
def vuln_scan_compare():
    """Compare two scans. Body: {scan_a: ref, scan_b: ref}
    Returns new findings (in B not A), fixed (in A not B), recurring (in both)."""
    data   = request.get_json(force=True)
    scan_a = _resolve_scan(data.get("scan_a", ""))
    scan_b = _resolve_scan(data.get("scan_b", ""))
    if not scan_a:
        return jsonify({"error": "scan_a not found"}), 404
    if not scan_b:
        return jsonify({"error": "scan_b not found"}), 404

    keys_a = {_finding_key(f): f for f in scan_a["findings"]}
    keys_b = {_finding_key(f): f for f in scan_b["findings"]}
    set_a, set_b = set(keys_a), set(keys_b)

    return jsonify({
        "scan_a": {k: scan_a[k] for k in ("name", "repo", "timestamp", "counts")},
        "scan_b": {k: scan_b[k] for k in ("name", "repo", "timestamp", "counts")},
        "new":       [keys_b[k] for k in set_b - set_a],
        "fixed":     [keys_a[k] for k in set_a - set_b],
        "recurring": [keys_b[k] for k in set_a & set_b],
    })


REPORT_MAX_CHARS = 60000

SEVERITY_ORDER = ["critical", "high", "medium", "low"]
SEVERITY_COLORS_HEX = {"critical": "4c0519", "high": "dc2626", "medium": "b45309", "low": "15803d"}


def count_findings_by_severity(findings):
    counts = {"critical": 0, "high": 0, "medium": 0, "low": 0}
    for f in findings:
        sev = (f.get("severity") or "").lower()
        if sev in counts:
            counts[sev] += 1
    return counts


DETAIL_BATCH_SIZE = 5      # findings per LLM call - small enough to stay specific
DETAIL_MAX_FINDINGS = 0    # 0 = no cap: EVERY finding gets a write-up. Batches
                           # run in parallel, so covering them all costs roughly
                           # (batches / DETAIL_PARALLEL) rounds rather than the
                           # sum of every call.
DETAIL_PARALLEL = 8        # batches in flight at once. The batches are fully
                           # independent, so running them sequentially made a
                           # large report take the SUM of every call instead of
                           # roughly the slowest one. Capped to stay polite to
                           # the LLM endpoint's rate limits.
DETAIL_RETRIES = 2         # a dropped batch means findings silently lose their write-up


def enrich_findings_with_detail(findings, log=None):
    """Asks the LLM to write a proper explanation for each finding: what it is,
    why it matters here, how it would be exploited, the specific fix, and how to
    verify it. Returns (details_by_index, coverage_note).

    Batched rather than one call per finding - a batch keeps each explanation
    grounded while cutting round trips by ~5x.

    Two details matter for coverage. Findings are selected in SEVERITY order,
    not scanner order, so if the cap is ever hit it drops the least important
    ones rather than whichever happened to be scanned last. And any finding
    missing from a batch's reply is retried, since a single malformed response
    would otherwise leave several findings silently unexplained."""
    client, model = get_client()
    if not client or not findings:
        return {}, ""

    order = {s: i for i, s in enumerate(SEVERITY_ORDER)}
    # Detections the AI already judged false positives go to an appendix as a
    # one-line row, so writing a six-paragraph analysis of each is pure waste -
    # on the last dashboard scan that was 100+ of 178 findings, and it made the
    # report both slower to produce and far harder to read.
    candidates = [i for i in range(len(findings)) if not _is_dismissed(findings[i])]
    ranked = sorted(
        candidates,
        key=lambda i: order.get((findings[i].get("severity") or "").lower(), len(SEVERITY_ORDER)),
    )
    # DETAIL_MAX_FINDINGS = 0 means no cap - every finding gets expanded.
    # Ordering by severity still matters: it decides which write-ups land
    # first, so the important ones are done even if a later batch fails.
    selected = ranked if DETAIL_MAX_FINDINGS <= 0 else ranked[:DETAIL_MAX_FINDINGS]

    def payload_for(idx_list):
        out = []
        for idx in idx_list:
            f = findings[idx]
            out.append({
                "id": str(idx),
                "file": f.get("file"),
                "line": f.get("line"),
                "severity": f.get("severity"),
                "category": f.get("category") or f.get("rule"),
                "description": f.get("description"),
                "evidence": (f.get("evidence") or "")[:600],
                "impact": f.get("impact"),
                "suggested_fix": f.get("improvement"),
                "scan_type": f.get("scan_type"),
            })
        return out

    def run_batch(chunk):
        """Returns {finding_index: detail} for one batch. Kept free of shared
        state so batches can run concurrently without locking."""
        out = {}
        raw = call_llm(client, model,
                       json.dumps(payload_for(chunk), indent=2)[:REPORT_MAX_CHARS],
                       system_prompt=VULN_DETAIL_SYSTEM_PROMPT)
        parsed = parse_json_response(raw)
        if isinstance(parsed, dict):
            parsed = [parsed]
        for item in parsed or []:
            if not isinstance(item, dict):
                continue
            try:
                idx = int(str(item.get("id")))
            except (TypeError, ValueError):
                continue
            # Only accept ids we actually asked about in this batch.
            if idx in chunk and any((item.get(k) or "").strip()
                                    for k in ("what", "why", "attack", "fix")):
                out[idx] = item
        return out

    details = {}
    pending = list(selected)
    for attempt in range(1, DETAIL_RETRIES + 2):
        if not pending:
            break
        if attempt > 1 and log is not None:
            log(f"Retrying detailed analysis for {len(pending)} finding(s) that came back empty...")
        batches = [pending[i:i + DETAIL_BATCH_SIZE] for i in range(0, len(pending), DETAIL_BATCH_SIZE)]
        if log is not None and attempt == 1:
            log(f"Writing detailed analysis: {len(batches)} batch(es), "
                f"up to {DETAIL_PARALLEL} at a time...")
        done = 0
        with ThreadPoolExecutor(max_workers=DETAIL_PARALLEL) as ex:
            futures = {ex.submit(run_batch, chunk): chunk for chunk in batches}
            for fut in as_completed(futures):
                done += 1
                try:
                    details.update(fut.result())
                except Exception as e:
                    if log is not None:
                        log(f"  detailed analysis failed for a batch: {e}")
                if log is not None and attempt == 1:
                    log(f"  detailed analysis: {done}/{len(batches)} batch(es) done")
        pending = [i for i in selected if i not in details]

    # Dismissed detections are intentionally not expanded, so they must not be
    # reported as a coverage gap - only a real cap counts as "skipped".
    skipped = len(candidates) - len(selected)
    note = ""
    if pending or skipped:
        bits = []
        if pending:
            bits.append(f"{len(pending)} finding(s) could not be expanded (the AI call failed "
                        "or returned nothing usable); their scanner-provided impact and fix are "
                        "shown instead")
        if skipped:
            bits.append(f"{skipped} lowest-severity finding(s) were not expanded, as this report "
                        f"caps detailed write-ups at {DETAIL_MAX_FINDINGS}")
        note = "; ".join(bits) + "."
    if log is not None:
        log(f"Detailed analysis written for {len(details)}/{len(selected)} finding(s).")
    return details, note


def _flatten_detail_value(val):
    """Collapses a detail field to a single clean line.

    The model sometimes returns a fenced code block despite being told not
    to. Fences don't survive this report's line-based markdown parser: the
    ``` markers leaked onto the page as stray backticks and the code lost
    its formatting anyway. Unwrapping them to inline code keeps the content
    without the broken rendering.
    """
    text = str(val or "").strip()
    if not text:
        return ""
    # ```lang\ncode\n``` -> `code`
    text = re.sub(r"```[a-zA-Z0-9_+-]*\s*\n?(.*?)```", lambda m: "`" + " ".join(m.group(1).split()) + "`",
                  text, flags=re.DOTALL)
    text = text.replace("```", "")
    return " ".join(text.split())


def _detail_section(detail):
    """Renders one finding's AI-written detail as markdown bullets."""
    if not detail:
        return []
    out = []
    # Deliberately short: what / why / how to fix. "verify" and the severity
    # commentary were dropped - with every finding expanded they roughly
    # doubled the length of each entry without changing what a developer
    # actually does about it. severity_note still surfaces below, but only
    # when the model flags the rating as wrong.
    for key, label in (("what", "What this is"), ("why", "Why it matters"),
                       ("attack", "How it's exploited"), ("fix", "Fix")):
        val = _flatten_detail_value(detail.get(key))
        if val:
            out.append(f"- **{label}:** {val}")
    note = _flatten_detail_value(detail.get("severity_note"))
    if note:
        out.append(f"- **Severity note:** {note}")
    return out


def _is_dismissed(f):
    """True when the AI review judged this finding not to be a real issue, or
    it was matched inside an obvious test fixture.

    Deliberately conservative: only an explicit false-positive verdict counts.
    An unreviewed finding, or one the AI confirmed, stays in the main report."""
    verdict = str(f.get("ai_verdict") or "").strip().lower()
    if verdict in ("false_positive", "false-positive", "falsepositive", "not_a_finding", "dismissed"):
        return True
    return bool(f.get("likely_test_fixture"))


def build_findings_markdown(result, details=None, detail_note=""):
    """Deterministically builds the Findings Summary + Detailed Findings
    sections directly from the scanner's JSON, so every single finding is
    guaranteed to appear in the report - not left to the LLM to enumerate,
    which risks silently dropping findings from a long list. Findings are
    grouped by severity (critical first) with a clear subsection per
    severity, so the document reads as an organized triage list rather
    than a flat dump."""
    all_findings = result.get("findings", [])

    # Split off what the AI review already judged not to be a real problem.
    # Mixing these in with genuine findings is what made a scan of 27 files
    # produce a 64-page report whose first three "critical" entries were all
    # test fixtures. They still belong in the document - a dismissal is a
    # judgement call, not proof - but as a separate, compact appendix rather
    # than at the top competing with real work.
    findings = [f for f in all_findings if not _is_dismissed(f)]
    dismissed = [f for f in all_findings if _is_dismissed(f)]

    counts = count_findings_by_severity(findings)

    category_counts = {}
    for f in findings:
        cat = f.get("category") or f.get("rule") or "uncategorized"
        category_counts[cat] = category_counts.get(cat, 0) + 1

    total = len(findings)

    def pct(n):
        return f"{round(100 * n / total)}%" if total else "0%"

    def loc_of(f):
        """file:line, but web/network findings have no line number - showing
        a bare ':None' there just looks like a bug in the report."""
        line = f.get("line")
        base = f.get("file", "unknown file")
        return f"{base}:{line}" if line not in (None, "") else base

    lines = ["## Findings Summary", ""]

    # --- Scope -----------------------------------------------------------
    scan_types = result.get("scan_types") or []
    affected_files = {f.get("file") for f in findings if f.get("file")}
    ai_found = sum(1 for f in findings if str(f.get("source") or "").startswith("ai-"))

    # --- Verdict ---------------------------------------------------------
    # A one-line risk posture up front, so the reader knows how bad it is
    # before working through the tables underneath.
    worst_sev = next((s for s in SEVERITY_ORDER if counts[s]), None)
    if not total:
        verdict, posture = "No issues found", "No findings were reported by this scan."
    elif worst_sev == "critical":
        verdict, posture = "Critical risk", "Critical issues are present and should be fixed before release."
    elif worst_sev == "high":
        verdict, posture = "High risk", "High-severity issues are present and should be prioritized."
    elif worst_sev == "medium":
        verdict, posture = "Moderate risk", "No critical or high issues; medium findings are worth scheduling."
    else:
        verdict, posture = "Low risk", "Only low-severity findings were reported."

    lines.append(f"> **Verdict: {verdict}.** {posture}")
    lines.append("")

    if total:
        # Lead with the single most urgent item so remediation has an obvious
        # starting point without reading the whole detail section.
        top = sorted(
            findings,
            key=lambda f: SEVERITY_ORDER.index((f.get("severity") or "").lower())
            if (f.get("severity") or "").lower() in SEVERITY_ORDER else len(SEVERITY_ORDER),
        )[0]
        top_loc = loc_of(top)
        lines.append(f"**Start here —** `{top_loc}` — "
                      f"{(top.get('severity') or 'unknown').capitalize()}: "
                      f"{top.get('description') or top.get('category') or 'see details below'}")
        lines.append("")

    lines.append("### Scope")
    lines.append("")
    lines.append("| | |")
    lines.append("|---|---|")
    lines.append(f"| Files scanned | {result.get('files_scanned', 'n/a')} |")
    lines.append(f"| Files with findings | {len(affected_files)} |")
    lines.append(f"| Raw detections | {len(all_findings)} |")
    lines.append(f"| **Findings needing review** | **{total}** |")
    if dismissed:
        lines.append(f"| Dismissed as likely false positives | {len(dismissed)} "
                     "(see appendix) |")
    if scan_types:
        lines.append(f"| Scan types | {', '.join(scan_types)} |")
    if ai_found:
        lines.append(f"| Found by AI review | {ai_found} of {len(all_findings)} |")
    if result.get("exit_code") is not None:
        lines.append(f"| Scanner exit code | {result.get('exit_code')} |")
    lines.append("")

    # --- Severity --------------------------------------------------------
    lines.append("### By severity")
    lines.append("")
    lines.append("| Severity | Count | Share | Distribution |")
    lines.append("|---|---|---|---|")
    for sev in SEVERITY_ORDER:
        n = counts[sev]
        # A text bar keeps the shape visible in markdown and Word too, not
        # only in the PDF where it can be drawn.
        bar = "█" * round(20 * n / total) if total and n else ""
        lines.append(f"| {sev.capitalize()} | {n} | {pct(n)} | {bar} |")
    lines.append(f"| **Total** | **{total}** | | |")
    lines.append("")

    # --- Categories ------------------------------------------------------
    if category_counts:
        # Worst severity per category makes the list triageable rather than a
        # bare frequency count - one critical outranks nine lows.
        worst = {}
        for f in findings:
            cat = f.get("category") or f.get("rule") or "uncategorized"
            sev = (f.get("severity") or "").lower()
            rank = SEVERITY_ORDER.index(sev) if sev in SEVERITY_ORDER else len(SEVERITY_ORDER)
            if cat not in worst or rank < worst[cat][0]:
                worst[cat] = (rank, sev or "unknown")

        lines.append("### By category")
        lines.append("")
        lines.append("| Category | Findings | Worst severity |")
        lines.append("|---|---|---|")
        for cat, n in sorted(category_counts.items(),
                             key=lambda kv: (worst.get(kv[0], (9, ""))[0], -kv[1])):
            lines.append(f"| {cat} | {n} | {worst.get(cat, (9, 'unknown'))[1].capitalize()} |")
        lines.append("")

    # --- Hotspots --------------------------------------------------------
    # Which files carry the most risk, so remediation can start where it pays
    # off most rather than at the top of an alphabetical list.
    file_stats = {}
    for f in findings:
        key = f.get("file") or "unknown"
        st = file_stats.setdefault(key, {"n": 0, "rank": len(SEVERITY_ORDER), "worst": "unknown"})
        st["n"] += 1
        sev = (f.get("severity") or "").lower()
        rank = SEVERITY_ORDER.index(sev) if sev in SEVERITY_ORDER else len(SEVERITY_ORDER)
        if rank < st["rank"]:
            st["rank"], st["worst"] = rank, sev or "unknown"
    if len(file_stats) > 1:
        top = sorted(file_stats.items(), key=lambda kv: (kv[1]["rank"], -kv[1]["n"]))[:10]
        lines.append("### Most affected files")
        lines.append("")
        lines.append("| File | Findings | Worst severity |")
        lines.append("|---|---|---|")
        for name, st in top:
            lines.append(f"| {name} | {st['n']} | {st['worst'].capitalize()} |")
        if len(file_stats) > len(top):
            lines.append(f"| _+{len(file_stats) - len(top)} more file(s)_ | | |")
        lines.append("")

    # --- Notes -----------------------------------------------------------
    skipped = result.get("skipped_files") or []
    ai_errors = result.get("ai_scan_errors") or []
    if result.get("truncated") or skipped or ai_errors or detail_note:
        lines.append("### Scan notes")
        lines.append("")
        if result.get("truncated"):
            lines.append("- Scan output was truncated; some files may not have been fully analyzed.")
        if skipped:
            lines.append(f"- {len(skipped)} file(s) were skipped during the scan (unsupported type, too large, or unreadable).")
        if ai_errors:
            lines.append(f"- AI verification failed for {len(ai_errors)} file(s); their findings are unverified.")
        if detail_note:
            lines.append(f"- {detail_note}")
        lines.append("")

    def sort_key(f):
        sev = (f.get("severity") or "low").lower()
        return SEVERITY_ORDER.index(sev) if sev in SEVERITY_ORDER else len(SEVERITY_ORDER)

    grouped = {sev: [] for sev in SEVERITY_ORDER}
    other = []
    # Detail keys are positions in the original findings list, so remember each
    # finding's index before regrouping by severity.
    index_of = {id(f): i for i, f in enumerate(findings)}
    for f in findings:
        sev = (f.get("severity") or "").lower()
        (grouped[sev] if sev in grouped else other).append(f)

    # Contents: an index of every finding by number, so a 20-page report can be
    # navigated without scrolling through it. Built from the same ordering the
    # detail sections use, so the numbers line up exactly.
    # No "Contents" index: it restated every finding a second time (about ten
    # pages on a real scan) while the severity/category/hotspot tables above
    # already give the overview, and each finding is numbered where it appears.
    lines.append("## Detailed Findings")
    lines.append("")

    counter = 0
    for sev in SEVERITY_ORDER + (["other"] if other else []):
        group = other if sev == "other" else grouped[sev]
        if not group:
            continue
        label = "Other" if sev == "other" else sev.capitalize()
        lines.append(f"### {label} severity ({len(group)})")
        lines.append("")
        for f in group:
            counter += 1
            sev_tag = (f.get("severity") or "unknown").upper()
            file_ = f.get("file", "unknown file")
            line_no = f.get("line", "?")
            category = f.get("category") or f.get("rule")
            # Numbered so findings can be referenced in review ("see #14")
            # rather than only by file path.
            heading = f"#### [{sev_tag}] #{counter} · {loc_of(f)}"
            if category:
                heading += f" — {category}"
            lines.append(heading)
            lines.append("")
            location = f"`{file_}`" + (f", line {line_no}" if line_no not in (None, "") else "")
            lines.append(f"- **Location:** {location}")
            if f.get("description"):
                lines.append(f"- **Issue:** {f['description']}")
            if f.get("evidence"):
                lines.append(f"- **Evidence:** `{f['evidence']}`")

            detail = (details or {}).get(index_of.get(id(f)))
            if detail:
                # The AI write-up supersedes the scanner's one-line impact/fix,
                # so those aren't repeated underneath it.
                lines.extend(_detail_section(detail))
            else:
                if f.get("impact"):
                    lines.append(f"- **Impact:** {f['impact']}")
                if f.get("improvement"):
                    lines.append(f"- **Fix:** {f['improvement']}")

            if f.get("ai_verdict"):
                verdict_line = f"- **AI verdict:** {f['ai_verdict']}"
                if f.get("ai_reason"):
                    verdict_line += f" ({f['ai_reason']})"
                lines.append(verdict_line)
            if f.get("likely_test_fixture"):
                lines.append("- _Likely a test fixture, not production code._")
            lines.append("")

    # --- Appendix: dismissed detections ----------------------------------
    # Listed compactly, one row each. These were judged not to be real issues,
    # so a full write-up per entry would bury the report in exactly the noise
    # this section exists to move out of the way - but they stay visible,
    # because a dismissal is a judgement call and can be wrong.
    if dismissed:
        lines.append("## Appendix — Possible False Positives")
        lines.append("")
        lines.append(f"{len(dismissed)} detection(s) were dismissed by AI review or matched "
                     "inside test fixtures, and are excluded from the counts above. They are "
                     "listed here rather than dropped: this is a best guess, not proof, so "
                     "scan the list before trusting it.")
        lines.append("")

        by_file = {}
        for f in dismissed:
            by_file.setdefault(f.get("file") or "unknown", []).append(f)

        def worst_rank(group):
            return min(
                (SEVERITY_ORDER.index((g.get("severity") or "").lower())
                 if (g.get("severity") or "").lower() in SEVERITY_ORDER else len(SEVERITY_ORDER))
                for g in group
            )

        for fname, group in sorted(by_file.items(), key=lambda kv: (worst_rank(kv[1]), -len(kv[1]))):
            group = sorted(
                group,
                key=lambda g: SEVERITY_ORDER.index((g.get("severity") or "").lower())
                if (g.get("severity") or "").lower() in SEVERITY_ORDER else len(SEVERITY_ORDER),
            )
            lines.append(f"### {fname} ({len(group)} dismissed)")
            lines.append("")
            lines.append("| Line | Severity | Category | Why it was dismissed |")
            lines.append("|---|---|---|---|")
            for f in group:
                line_no = f.get("line")
                reason = (f.get("ai_reason") or
                          ("Matched inside a test fixture" if f.get("likely_test_fixture")
                           else "Dismissed by AI review"))
                reason = str(reason).replace("|", "\\|")
                cat = (f.get("category") or f.get("rule") or "uncategorized").replace("|", "\\|")
                lines.append(f"| {line_no if line_no not in (None, '') else '—'} "
                             f"| {(f.get('severity') or 'unknown').capitalize()} "
                             f"| {cat} | {reason} |")
            lines.append("")

    return "\n".join(lines)


def parse_markdown_blocks(md_text):
    """Small markdown parser tuned to the structure this report's sections
    produce (## / ### headings, "- " bullets, "1. " numbered lists,
    paragraphs). Returns a list of (block_type, text) tuples for the
    docx/pdf builders."""
    blocks = []
    table_rows = None   # collects consecutive "| a | b |" lines

    def flush_table():
        nonlocal table_rows
        if table_rows:
            blocks.append(("table", table_rows))
            table_rows = None

    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()

        # Pipe tables: gather the run of rows, dropping the |---|---| separator
        # and any fully empty header used only to make a two-column layout.
        if line.strip().startswith("|") and line.strip().endswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if not all(re.fullmatch(r":?-{2,}:?", c or "-") for c in cells):
                if table_rows is None:
                    table_rows = []
                table_rows.append(cells)
            continue
        flush_table()

        if not line.strip():
            continue
        if line.startswith("#### "):
            blocks.append(("h4", line[5:].strip()))
        elif line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
        elif line.lstrip().startswith("- "):
            blocks.append(("bullet", line.lstrip()[2:].strip()))
        elif re.match(r"^\d+\.\s", line.lstrip()):
            blocks.append(("numbered", re.sub(r"^\d+\.\s", "", line.lstrip())))
        else:
            blocks.append(("p", line.strip()))
    return blocks


# Single-asterisk emphasis, matched only when it isn't part of a ** pair.
# Applied after bold so "**x**" is already consumed by then.
_MD_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\s)([^*\n]+?)(?<!\s)\*(?!\*)")


def strip_md_inline(text):
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = _MD_ITALIC_RE.sub(r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


FINDING_HEADING_RE = re.compile(r"^\[(CRITICAL|HIGH|MEDIUM|LOW|UNKNOWN)\]\s*(.*)$")


def build_report_docx(md_text, title, subtitle=None):
    from docx import Document
    from docx.shared import RGBColor

    doc = Document()
    doc.add_heading(title, level=0)
    if subtitle:
        sub = doc.add_paragraph(subtitle)
        sub.runs[0].italic = True

    for kind, text in parse_markdown_blocks(md_text):
        # Tables arrive as a list of row-cell-lists, so they must be handled
        # before any of the string-only formatting below.
        if kind == "table":
            rows = [r for r in text if any(c.strip() for c in r)]
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            tbl = doc.add_table(rows=0, cols=ncols)
            tbl.style = "Light Grid Accent 1"
            header_is_blank = not any(c.strip() for c in rows[0])
            for i, row in enumerate(rows):
                cells = tbl.add_row().cells
                for j in range(ncols):
                    val = strip_md_inline(row[j]) if j < len(row) else ""
                    cells[j].text = val
                    if i == 0 and not header_is_blank:
                        for para in cells[j].paragraphs:
                            for run in para.runs:
                                run.bold = True
            doc.add_paragraph()
            continue
        if kind == "h4":
            m = FINDING_HEADING_RE.match(text)
            p = doc.add_paragraph(style="Heading 4")
            if m:
                sev, rest = m.group(1), strip_md_inline(m.group(2))
                hexcolor = SEVERITY_COLORS_HEX.get(sev.lower(), "374151")
                run = p.add_run(sev + "  ")
                run.font.color.rgb = RGBColor.from_string(hexcolor)
                p.add_run(rest)
            else:
                p.add_run(strip_md_inline(text))
            continue
        text = strip_md_inline(text)
        if kind == "h1":
            doc.add_heading(text, level=1)
        elif kind == "h2":
            doc.add_heading(text, level=2)
        elif kind == "h3":
            doc.add_heading(text, level=3)
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "numbered":
            doc.add_paragraph(text, style="List Number")
        else:
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_REPORT_FONTS = None


def _register_report_fonts():
    """Registers a real Unicode TrueType font for the PDF, falling back through
    a few candidates.

    The standard-14 PDF fonts (Helvetica et al.) are Type1 with WinAnsi
    encoding, which simply has no glyph for characters like the em-dash in a
    finding heading, an arrow, or a check mark - they come out blank or as a
    box. A TrueType font embedded with UTF-8 encoding renders all of them.

    Returns (regular, bold, mono) font names; falls back to the built-ins if
    nothing can be registered, in which case _ascii_fallback strips what
    wouldn't render."""
    global _REPORT_FONTS
    if _REPORT_FONTS is not None:
        return _REPORT_FONTS

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import reportlab as _rl

    rl_fonts = Path(_rl.__file__).parent / "fonts"
    win_fonts = Path(os.environ.get("SYSTEMROOT", r"C:\Windows")) / "Fonts"
    candidates = [
        # (family, regular, bold) - best Unicode coverage first.
        ("DejaVu", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                   "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        ("Arial", win_fonts / "arial.ttf", win_fonts / "arialbd.ttf"),
        ("Vera", rl_fonts / "Vera.ttf", rl_fonts / "VeraBd.ttf"),
    ]
    mono_candidates = [
        ("DejaVuMono", "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf"),
        ("Consolas", win_fonts / "consola.ttf"),
        ("CourierNew", win_fonts / "cour.ttf"),
    ]

    regular, bold, mono = "Helvetica", "Helvetica-Bold", "Courier"
    for name, reg_path, bold_path in candidates:
        try:
            if not (Path(reg_path).exists() and Path(bold_path).exists()):
                continue
            pdfmetrics.registerFont(TTFont(name, str(reg_path)))
            pdfmetrics.registerFont(TTFont(name + "-Bold", str(bold_path)))
            regular, bold = name, name + "-Bold"
            break
        except Exception:
            continue
    for name, path in mono_candidates:
        try:
            if not Path(path).exists():
                continue
            pdfmetrics.registerFont(TTFont(name, str(path)))
            mono = name
            break
        except Exception:
            continue

    _REPORT_FONTS = (regular, bold, mono)
    return _REPORT_FONTS


# Only used when no TrueType font could be registered: map characters the
# built-in WinAnsi fonts can't draw onto plain ASCII, so they degrade to
# something readable instead of vanishing.
_ASCII_FALLBACKS = {
    "\u2014": "-", "\u2013": "-", "\u2212": "-",
    "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"',
    "\u2026": "...", "\u00b7": "-", "\u2022": "*",
    "\u2192": "->", "\u2190": "<-", "\u2191": "^", "\u2193": "v",
    "\u2713": "OK", "\u2717": "x", "\u00a0": " ",
    "\u2588": "#",
}

# Characters that are visually near-identical to a plain ASCII one but are
# frequently MISSING from a given font, so they render as a tofu box instead.
# LLM prose is full of these (U+2011 in "hard-coded", narrow/thin spaces),
# and they carry no meaning worth the rendering risk - so they're normalized
# away for every output format, not only the ASCII-only PDF fallback path.
_RISKY_CHAR_NORMALIZATIONS = {
    "\u2010": "-",   # hyphen
    "\u2011": "-",   # non-breaking hyphen  <- the usual "box" culprit
    "\u2012": "-",   # figure dash
    "\u2015": "-",   # horizontal bar
    "\u00ad": "",    # soft hyphen (invisible, but can render as a box)
    "\u2009": " ",   # thin space
    "\u202f": " ",   # narrow no-break space
    "\u200b": "",    # zero-width space
    "\u200e": "", "\u200f": "",  # bidi marks
    "\ufeff": "",    # BOM / zero-width no-break space
}


def normalize_report_text(text):
    if not text:
        return text
    for bad, good in _RISKY_CHAR_NORMALIZATIONS.items():
        text = text.replace(bad, good)
    return text


def _ascii_fallback(text, active):
    if not active or not text:
        return text
    for bad, good in _ASCII_FALLBACKS.items():
        text = text.replace(bad, good)
    return text


def build_report_pdf(md_text, title, subtitle=None, meta=None):
    import html as _html
    from xml.sax.saxutils import escape as xml_escape
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor, white
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak,
        Table, TableStyle, HRFlowable, KeepTogether, CondPageBreak,
    )

    W, H = letter
    buf = io.BytesIO()
    F_REG, F_BOLD, F_MONO = _register_report_fonts()
    # True only if no TrueType font could be registered, in which case
    # characters outside WinAnsi have to be transliterated to survive.
    ASCII_ONLY = F_REG == "Helvetica"

    # Palette
    C_NAVY   = HexColor("#1e293b")
    C_SLATE  = HexColor("#475569")
    C_MUTED  = HexColor("#94a3b8")
    C_BORDER = HexColor("#e2e8f0")

    SEV_FG = {
        "critical": HexColor("#dc2626"),
        "high":     HexColor("#ea580c"),
        "medium":   HexColor("#d97706"),
        "low":      HexColor("#16a34a"),
        "unknown":  HexColor("#6b7280"),
    }
    SEV_BG = {
        "critical": HexColor("#fef2f2"),
        "high":     HexColor("#fff7ed"),
        "medium":   HexColor("#fffbeb"),
        "low":      HexColor("#f0fdf4"),
        "unknown":  HexColor("#f9fafb"),
    }
    SEV_HDR_BG = {
        "critical": HexColor("#7f1d1d"),
        "high":     HexColor("#7c2d12"),
        "medium":   HexColor("#78350f"),
        "low":      HexColor("#14532d"),
        "unknown":  HexColor("#374151"),
    }

    # ── Canvas callbacks ────────────────────────────────────────────────────
    def _cover(c, doc):
        # Dark navy background
        c.setFillColor(C_NAVY)
        c.rect(0, 0, W, H, fill=1, stroke=0)
        # Blue accent stripe
        c.setFillColor(HexColor("#3b82f6"))
        c.rect(0.75 * inch, H * 0.415, W - 1.5 * inch, 3, fill=1, stroke=0)
        # Eyebrow
        c.setFillColor(HexColor("#64748b"))
        c.setFont(F_REG, 10)
        c.drawCentredString(W / 2, H * 0.72, "SECURITY  ·  CONFIDENTIAL")
        # Big title
        c.setFillColor(white)
        c.setFont(F_BOLD, 38)
        c.drawCentredString(W / 2, H * 0.615, "VULNERABILITY")
        c.setFont(F_BOLD, 30)
        c.drawCentredString(W / 2, H * 0.545, "ASSESSMENT REPORT")
        # Subtitle (date / files / count line)
        c.setFillColor(HexColor("#94a3b8"))
        c.setFont(F_REG, 11)
        sub = _ascii_fallback((subtitle or "").replace("&middot;", "·"), ASCII_ONLY)
        c.drawCentredString(W / 2, H * 0.43, sub)
        # Bottom stats panel
        counts = (meta or {}).get("counts", {})
        files  = (meta or {}).get("files_scanned", "?")
        total  = sum(counts.values()) if counts else "?"
        c.setFillColor(HexColor("#0f172a"))
        c.rect(0, 0, W, H * 0.29, fill=1, stroke=0)
        c.setFillColor(HexColor("#334155"))
        c.rect(0, H * 0.29, W, 1, fill=1, stroke=0)
        stats = [
            (str(files),                        "FILES SCANNED",  "#94a3b8"),
            (str(total),                         "TOTAL FINDINGS", "#cbd5e1"),
            (str(counts.get("critical", 0)),     "CRITICAL",       "#dc2626"),
            (str(counts.get("high", 0)),         "HIGH",           "#ea580c"),
            (str(counts.get("medium", 0)),       "MEDIUM",         "#d97706"),
            (str(counts.get("low", 0)),          "LOW",            "#16a34a"),
        ]
        col_w = W / len(stats)
        for i, (val, lbl, clr) in enumerate(stats):
            cx = i * col_w + col_w / 2
            c.setFillColor(HexColor(clr))
            c.setFont(F_BOLD, 26)
            c.drawCentredString(cx, H * 0.165, val)
            c.setFillColor(HexColor("#64748b"))
            c.setFont(F_REG, 8)
            c.drawCentredString(cx, H * 0.105, lbl)

    def _page(c, doc):
        c.saveState()
        # Top bar
        c.setFillColor(C_NAVY)
        c.rect(0, H - 0.44 * inch, W, 0.44 * inch, fill=1, stroke=0)
        c.setFillColor(white)
        c.setFont(F_BOLD, 8)
        c.drawString(0.75 * inch, H - 0.27 * inch, title)
        c.setFont(F_REG, 8)
        c.drawRightString(W - 0.75 * inch, H - 0.27 * inch, "CONFIDENTIAL")
        # Footer
        c.setStrokeColor(C_BORDER)
        c.line(0.75 * inch, 0.54 * inch, W - 0.75 * inch, 0.54 * inch)
        c.setFillColor(C_MUTED)
        c.setFont(F_REG, 7.5)
        # The full subtitle carries an absolute path plus counts plus a
        # timestamp - repeating all of it on every page is noise. The cover
        # already states it in full, so the footer keeps just the target.
        sub = _ascii_fallback((subtitle or "").replace("&middot;", "·"), ASCII_ONLY)
        short_sub = sub.split(" - ")[0].strip()
        if len(short_sub) > 70:
            short_sub = "..." + short_sub[-67:]
        c.drawString(0.75 * inch, 0.36 * inch, short_sub)
        c.drawRightString(W - 0.75 * inch, 0.36 * inch, f"Page {doc.page}")
        c.restoreState()

    # ── Styles ──────────────────────────────────────────────────────────────
    PS = ParagraphStyle
    S_H1    = PS("RH1",    fontName=F_BOLD, fontSize=20, textColor=C_NAVY,
                 spaceBefore=20, spaceAfter=8, leading=24)
    S_H2    = PS("RH2",    fontName=F_BOLD, fontSize=15, textColor=C_NAVY,
                 spaceBefore=18, spaceAfter=4, leading=18)
    S_H3    = PS("RH3",    fontName=F_BOLD, fontSize=12, textColor=C_SLATE,
                 spaceBefore=10, spaceAfter=4, leading=15)
    S_H3W   = PS("RH3W",   fontName=F_BOLD, fontSize=11, textColor=white,
                 spaceBefore=0,  spaceAfter=0, leading=14, leftIndent=8)
    S_BODY  = PS("RBody",  fontName=F_REG,      fontSize=10,
                 textColor=HexColor("#374151"), spaceBefore=3, spaceAfter=3, leading=15)
    S_BULL  = PS("RBull",  fontName=F_REG,      fontSize=10,
                 textColor=HexColor("#374151"), spaceBefore=2, spaceAfter=2,
                 leftIndent=12, leading=14)
    S_BADGE = PS("RBadge", fontName=F_BOLD, fontSize=8.5, textColor=white,
                 alignment=TA_CENTER, spaceBefore=0, spaceAfter=0, leading=11)
    S_FTIT  = PS("RFTit",  fontName=F_BOLD, fontSize=9.5, textColor=white,
                 spaceBefore=0, spaceAfter=0, leading=13)
    S_FBOD  = PS("RFBod",  fontName=F_REG,      fontSize=9,
                 textColor=HexColor("#374151"), spaceBefore=2, spaceAfter=2, leading=13)
    S_FCOD  = PS("RFCod",  fontName=F_MONO,        fontSize=8,
                 textColor=HexColor("#1e293b"), spaceBefore=2, spaceAfter=2,
                 backColor=HexColor("#f1f5f9"), leftIndent=4, leading=11)

    S_TH    = PS("RTH",    fontName=F_BOLD, fontSize=8.5,
                 textColor=HexColor("#334155"), spaceBefore=0, spaceAfter=0, leading=11)
    S_TD    = PS("RTD",    fontName=F_REG,  fontSize=8.5,
                 textColor=HexColor("#374151"), spaceBefore=0, spaceAfter=0, leading=11)

    C_ACCENT = HexColor("#4f46e5")
    CNTW = W - 1.5 * inch  # usable content width

    def _safe(text):
        """Unescape HTML entities then re-escape for reportlab XML."""
        return xml_escape(_ascii_fallback(_html.unescape(strip_md_inline(text)), ASCII_ONLY))

    def _rich(text):
        """Like _safe, but keeps **bold** and `code` as real formatting instead
        of stripping the markers. Escaping happens first, so nothing in the
        source text can inject reportlab markup."""
        out = xml_escape(_ascii_fallback(_html.unescape(text or ""), ASCII_ONLY))
        out = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", out)
        # Without this, the model's *emphasis* reached the page as literal
        # asterisks ("*Why it matters:*").
        out = _MD_ITALIC_RE.sub(r"<i>\1</i>", out)
        out = re.sub(r"`(.+?)`", rf'<font face="{F_MONO}" size="8.5">\1</font>', out)
        return out

    def _md_table(rows):
        """Renders a markdown pipe-table as a real styled PDF table. A leading
        all-blank header row means the source used a headerless two-column
        layout for key/value pairs, so it's dropped and the first column is
        styled as labels instead of drawing an empty header band."""
        rows = [r for r in rows if any(c.strip() for c in r)]
        if not rows:
            return []
        ncols = max(len(r) for r in rows)
        rows = [list(r) + [""] * (ncols - len(r)) for r in rows]
        has_header = any(c.strip() for c in rows[0]) and len(rows) > 1

        # A key/value table wants a narrow label column, but a two-column
        # index of findings wants even halves - decide from how much text the
        # first column actually carries rather than assuming it's a label.
        body_rows = rows[1:] if has_header else rows
        avg_first = (sum(len(r[0]) for r in body_rows) / len(body_rows)) if body_rows else 0
        if avg_first > 25:
            col_widths = [CNTW / ncols] * ncols
            rest_w = CNTW / ncols
        else:
            first_w = min(2.9 * inch, CNTW * 0.42)
            rest_w = (CNTW - first_w) / max(1, ncols - 1) if ncols > 1 else 0
            col_widths = [first_w] + [rest_w] * (ncols - 1)

        def _bar(units):
            """A run of block characters is a text bar chart in the source; in
            the PDF it's drawn as a real colored bar, so it doesn't depend on
            the chosen font actually having that glyph."""
            frac = max(0.03, min(1.0, units / 20))
            bar = Table([[""]], colWidths=[max(2, (rest_w - 16) * frac)], rowHeights=[7], hAlign="LEFT")
            bar.setStyle(TableStyle([
                ("BACKGROUND",   (0, 0), (-1, -1), C_ACCENT),
                ("LEFTPADDING",  (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING",   (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]))
            return bar

        data = []
        for i, row in enumerate(rows):
            cells = []
            for cell in row:
                stripped = cell.strip()
                if stripped and set(stripped) == {"█"}:
                    cells.append(_bar(len(stripped)))
                else:
                    style = S_TH if (i == 0 and has_header) else S_TD
                    cells.append(Paragraph(_rich(cell), style))
            data.append(cells)

        t = Table(data, colWidths=col_widths, hAlign="LEFT")
        style = [
            ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING",   (0, 0), (-1, -1), 8),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 8),
            ("TOPPADDING",    (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LINEBELOW",     (0, 0), (-1, -2), 0.5, C_BORDER),
            ("BOX",           (0, 0), (-1, -1), 0.7, C_BORDER),
        ]
        if has_header:
            style += [
                ("BACKGROUND", (0, 0), (-1, 0), HexColor("#f1f5f9")),
                ("LINEBELOW",  (0, 0), (-1, 0), 0.9, C_BORDER),
            ]
        for r in range(1 if has_header else 0, len(data)):
            if r % 2 == (1 if has_header else 0):
                style.append(("BACKGROUND", (0, r), (-1, r), HexColor("#fbfcfe")))
        t.setStyle(TableStyle(style))

        return [Spacer(1, 4), t, Spacer(1, 8)]

    def _h2_with_rule(text_esc):
        t = Table(
            [[Paragraph(text_esc, S_H2)], [HRFlowable(width=CNTW, color=C_BORDER, thickness=1)]],
            colWidths=[CNTW],
        )
        t.setStyle(TableStyle([("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                                ("TOPPADDING",    (0, 0), (-1, -1), 0)]))
        return t

    def _sev_group_header(text_esc, sev_key):
        bg = SEV_HDR_BG.get(sev_key, HexColor("#374151"))
        t = Table([[Paragraph(text_esc, S_H3W)]], colWidths=[CNTW])
        t.setStyle(TableStyle([
            ("BACKGROUND",     (0, 0), (-1, -1), bg),
            ("TOPPADDING",     (0, 0), (-1, -1), 9),
            ("BOTTOMPADDING",  (0, 0), (-1, -1), 9),
            ("LEFTPADDING",    (0, 0), (-1, -1), 10),
            ("RIGHTPADDING",   (0, 0), (-1, -1), 10),
        ]))
        # Demand room for the header plus the start of its first finding, so a
        # bar reading "Critical severity (10)" can never be stranded alone at
        # the foot of a page with its findings overleaf.
        return [CondPageBreak(2.6 * inch), Spacer(1, 10), t, Spacer(1, 4)]

    def _finding_card(h4_text, card_blocks):
        m = FINDING_HEADING_RE.match(h4_text)
        sev_key = m.group(1).lower() if m else "unknown"
        rest    = _html.unescape(strip_md_inline(m.group(2))) if m else _html.unescape(strip_md_inline(h4_text))
        fg  = SEV_FG.get(sev_key, HexColor("#6b7280"))
        bg  = SEV_BG.get(sev_key, HexColor("#f9fafb"))

        badge = Paragraph(f'<b>{sev_key.upper()}</b>', S_BADGE)
        titl  = Paragraph(xml_escape(rest), S_FTIT)
        hdr   = Table([[badge, titl]], colWidths=[0.68 * inch, CNTW - 0.68 * inch])
        hdr.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0), (0, -1), fg),
            ("BACKGROUND",    (1, 0), (1, -1), C_NAVY),
            ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING",    (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ("LEFTPADDING",   (0, 0), (0,  -1), 5),
            ("RIGHTPADDING",  (0, 0), (0,  -1), 5),
            ("LEFTPADDING",   (1, 0), (1,  -1), 10),
            ("RIGHTPADDING",  (1, 0), (1,  -1), 8),
        ]))

        # A Paragraph inside a table cell cannot be split by reportlab, so a
        # single very long value (a verbose AI write-up, a huge minified
        # evidence line) would be an unsplittable flowable taller than the
        # page and abort the whole document. Splitting the RAW text into
        # page-sized chunks first keeps every row small enough to lay out.
        CHUNK = 1400

        def _chunk_raw(raw):
            raw = raw or ""
            if len(raw) <= CHUNK:
                return [raw]
            out, rest = [], raw
            while len(rest) > CHUNK:
                cut = rest.rfind(" ", 0, CHUNK)
                if cut <= 0:
                    cut = CHUNK
                out.append(rest[:cut])
                rest = rest[cut:].lstrip()
            if rest:
                out.append(rest)
            return out

        body_paras = []
        for bkind, btext in card_blocks:
            # A pipe-table can land inside a card when AI-written prose happens
            # to contain "|" characters. btext is then a list of rows, so it
            # must be flattened before any string handling below.
            if bkind == "table" or not isinstance(btext, str):
                for row in (btext if isinstance(btext, list) else []):
                    cells = [c for c in row if str(c).strip()]
                    if cells:
                        body_paras.append(Paragraph(_rich(" | ".join(str(c) for c in cells)), S_FBOD))
                continue
            lbl_m = re.match(r"^\*\*(.+?):\*\*\s*(.*)", btext)
            if lbl_m and bkind == "bullet":
                lbl = xml_escape(lbl_m.group(1))
                # Evidence / code lines get monospace treatment
                if lbl_m.group(1).lower() in ("evidence", "code"):
                    body_paras.append(Paragraph(f'<b>{lbl}:</b>', S_FBOD))
                    raw_code = _ascii_fallback(_html.unescape(strip_md_inline(lbl_m.group(2))), ASCII_ONLY)
                    for piece in _chunk_raw(raw_code):
                        body_paras.append(Paragraph(xml_escape(piece), S_FCOD))
                else:
                    pieces = _chunk_raw(lbl_m.group(2))
                    body_paras.append(Paragraph(f'<b>{lbl}:</b>  {_rich(pieces[0])}', S_FBOD))
                    for piece in pieces[1:]:
                        body_paras.append(Paragraph(_rich(piece), S_FBOD))
            else:
                prefix = "• " if bkind == "bullet" else ""
                pieces = _chunk_raw(btext)
                body_paras.append(Paragraph(prefix + _rich(pieces[0]), S_FBOD))
                for piece in pieces[1:]:
                    body_paras.append(Paragraph(_rich(piece), S_FBOD))

        # The card's border is drawn per-section rather than by wrapping the
        # whole thing in an outer Table. A KeepTogether nested inside a table
        # cell has no computable height, which makes reportlab raise
        # LayoutError the moment a card lands near a page break.
        hdr.setStyle(TableStyle([
            ("LINEABOVE",  (0, 0), (-1, 0), 1, fg),
            ("LINEBEFORE", (0, 0), (0, -1), 1, fg),
            ("LINEAFTER",  (-1, 0), (-1, -1), 1, fg),
        ]))

        elements = [hdr]
        if body_paras:
            # One row PER paragraph, not one cell holding them all. reportlab
            # can only split a table between rows - a single row taller than
            # the page raises LayoutError and kills the whole document, which
            # is exactly what a long AI write-up produces.
            body_tbl = Table([[p] for p in body_paras], colWidths=[CNTW], repeatRows=0)
            body_tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0, 0), (-1, -1), bg),
                ("LEFTPADDING",   (0, 0), (-1, -1), 10),
                ("RIGHTPADDING",  (0, 0), (-1, -1), 10),
                ("TOPPADDING",    (0, 0), (0, 0), 8),
                ("BOTTOMPADDING", (0, -1), (-1, -1), 8),
                ("TOPPADDING",    (0, 1), (-1, -1), 1),
                ("BOTTOMPADDING", (0, 0), (-1, -2), 1),
                ("VALIGN",        (0, 0), (-1, -1), "TOP"),
                ("LINEBEFORE",    (0, 0), (0, -1), 1, fg),
                ("LINEAFTER",     (-1, 0), (-1, -1), 1, fg),
                ("LINEBELOW",     (0, -1), (-1, -1), 1, fg),
            ]))
            elements.append(body_tbl)
        else:
            hdr.setStyle(TableStyle([("LINEBELOW", (0, -1), (-1, -1), 1, fg)]))

        # Only keep the card intact when it's short enough to actually fit on
        # one page; forcing KeepTogether on a long card re-creates the same
        # unsplittable-flowable crash it's meant to avoid.
        if len(body_paras) <= 6:
            return [Spacer(1, 6), KeepTogether(elements)]
        return [Spacer(1, 6)] + elements

    # ── Pre-group blocks: flatten h4 + following bullets into finding cards ─
    raw_blocks = parse_markdown_blocks(md_text)
    grouped = []
    cur_h4, cur_card = None, []
    for kind, text in raw_blocks:
        if kind in ("h1", "h2", "h3"):
            if cur_h4 is not None:
                grouped.append(("finding_card", cur_h4, cur_card))
                cur_h4, cur_card = None, []
            grouped.append((kind, text))
        elif kind == "h4":
            if cur_h4 is not None:
                grouped.append(("finding_card", cur_h4, cur_card))
            cur_h4, cur_card = text, []
        else:
            if cur_h4 is not None:
                cur_card.append((kind, text))
            else:
                grouped.append((kind, text))
    if cur_h4 is not None:
        grouped.append(("finding_card", cur_h4, cur_card))

    # ── Build story ─────────────────────────────────────────────────────────
    story = [PageBreak()]  # cover is drawn by _cover; PageBreak starts page 2

    for item in grouped:
        kind = item[0]
        text = item[1] if len(item) > 1 else ""

        if kind == "finding_card":
            story.extend(_finding_card(item[1], item[2]))
            continue

        # Tables carry a list of rows rather than a string, so they must be
        # handled before anything that assumes text is a string.
        if kind == "table":
            story.extend(_md_table(text))
            continue

        text_esc = _safe(text)

        if kind == "h1":
            story.extend([CondPageBreak(1.8 * inch), Spacer(1, 10), Paragraph(text_esc, S_H1)])
        elif kind == "h2":
            # The contents index and the finding list each get their own page,
            # which keeps the summary sections intact and makes the document
            # easier to navigate.
            if text.strip().lower().startswith(("detailed findings", "contents")):
                story.append(PageBreak())
                story.append(_h2_with_rule(text_esc))
            else:
                story.extend([CondPageBreak(1.8 * inch), Spacer(1, 14), _h2_with_rule(text_esc)])
        elif kind == "h3":
            sev_m = re.match(
                r"^(Critical|High|Medium|Low|Other)\s+severity", text, re.IGNORECASE
            )
            if sev_m:
                story.extend(_sev_group_header(text_esc, sev_m.group(1).lower()))
            else:
                story.extend([CondPageBreak(1.4 * inch), Spacer(1, 8), Paragraph(text_esc, S_H3)])
        elif kind == "bullet":
            story.append(Paragraph(f"• {_rich(text)}", S_BULL))
        else:
            story.append(Paragraph(_rich(text), S_BODY))

        story.append(Spacer(1, 3))

    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        topMargin=0.85 * inch, bottomMargin=0.78 * inch,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
    )
    doc.build(story, onFirstPage=_cover, onLaterPages=_page)
    return buf.getvalue()


def build_report_html(md_text, title, subtitle=None):
    import html as html_lib
    import markdown as md_lib

    def colorize_heading(m):
        sev = m.group(1)
        hexcolor = SEVERITY_COLORS_HEX.get(sev.lower(), "374151")
        return (
            f'#### <span style="background:#{hexcolor}22;color:#{hexcolor};'
            f'padding:2px 10px;border-radius:999px;font-size:0.68em;'
            f'font-weight:700;letter-spacing:0.03em;">{sev}</span>'
        )

    md_text = re.sub(r"^#### \[(CRITICAL|HIGH|MEDIUM|LOW|UNKNOWN)\]", colorize_heading, md_text, flags=re.MULTILINE)
    body = md_lib.markdown(md_text, extensions=["extra"])
    safe_title = html_lib.escape(title)
    # subtitle is built server-side with a literal `&middot;` entity, not
    # user input, so it's passed through as-is rather than re-escaped
    # (escaping would turn it into the literal text "&middot;").
    safe_subtitle = subtitle or ""
    return (
        "<!doctype html><html><head><meta charset=\"utf-8\">"
        f"<title>{safe_title}</title><style>"
        "body{font-family:-apple-system,Segoe UI,Arial,sans-serif;max-width:860px;"
        "margin:2rem auto;padding:0 1.5rem;color:#1f2937;line-height:1.6;}"
        "h1{margin-bottom:2px;}"
        ".report-subtitle{color:#9ca3af;font-size:13px;margin-bottom:1.5rem;}"
        "h2{color:#111827;border-bottom:1px solid #e5e7eb;padding-bottom:4px;margin-top:2.2rem;}"
        "h3{margin-top:1.8rem;color:#374151;}"
        "h4{display:flex;align-items:center;gap:10px;margin-top:1.4rem;font-size:15px;}"
        "code{background:#f3f4f6;padding:1px 4px;border-radius:4px;}"
        "ul{padding-left:1.3rem;}"
        "table{border-collapse:collapse;width:100%;margin:0.75rem 0 1.25rem;font-size:13.5px;}"
        "th{background:#f1f5f9;color:#334155;text-align:left;font-weight:600;}"
        "th,td{border:1px solid #e2e8f0;padding:7px 10px;vertical-align:middle;}"
        "tbody tr:nth-child(even){background:#fbfcfe;}"
        # The distribution column is a run of block characters; coloring it
        # makes it read as a bar chart rather than a wall of dark glyphs.
        "td:last-child{color:#4f46e5;letter-spacing:-1px;}"
        "blockquote{margin:0 0 1rem;padding:10px 14px;background:#eef2ff;"
        "border-left:4px solid #4f46e5;border-radius:0 6px 6px 0;color:#312e81;}"
        "blockquote p{margin:0;}"
        f"</style></head><body><h1>{safe_title}</h1>"
        f"<div class=\"report-subtitle\">{safe_subtitle}</div>{body}</body></html>"
    )


def _has_weak_severity_note(detail_dict):
  """Returns True if the severity_note indicates uncertainty or low confidence."""
  note = (detail_dict or {}).get("severity_note", "").lower()
  weak_signals = (
    "false positive", "no risk", "no user data", "static", "empty string",
    "likely a false positive", "appears to be a false positive",
    "needs more context", "insufficient context", "further context",
    "no viable", "unclear", "cannot exploit", "safe as written", "overstated"
  )
  return any(signal in note for signal in weak_signals)


def build_fix_loop_markdown(result, details=None, scan_label="", scan_time=None):
    """Builds a task file meant to be handed straight to a coding agent.

    A report is written for a person deciding what to do. This is written for an
    agent doing it: one task per finding, ordered worst-first, each carrying the
    exact location, the offending line, what to change and how to prove it
    worked - plus a checkbox so a long run can be stopped and resumed.

    Two things make the difference between this and just pasting the report:
    dismissed detections are listed as explicitly DO-NOT-FIX (an agent handed
    178 findings will otherwise "fix" the AWS key in a redaction test and break
    it), and the rules up front stop it inventing changes for findings whose
    evidence is too thin to act on."""
    all_findings = result.get("findings", [])
    findings = [f for f in all_findings if not _is_dismissed(f)]
    dismissed = [f for f in all_findings if _is_dismissed(f)]
    details = details or {}

    order = {s: i for i, s in enumerate(SEVERITY_ORDER)}
    index_of = {id(f): i for i, f in enumerate(all_findings)}

    # Separate findings into actionable and uncertain (weak severity notes)
    actionable = []
    uncertain = []
    for f in findings:
        idx = index_of.get(id(f), 0)
        d = details.get(idx, {}) or {}
        if _has_weak_severity_note(d):
            uncertain.append(f)
        else:
            actionable.append(f)

    ranked = sorted(actionable, key=lambda f: order.get((f.get("severity") or "").lower(),
                                                        len(SEVERITY_ORDER)))
    counts = count_findings_by_severity(ranked)
    stamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    L = []
    L.append("# Security fix loop")
    L.append("")
    L.append(f"- **Target:** `{scan_label or 'scanned project'}`")
    L.append(f"- **Generated:** {stamp}")
    L.append(f"- **Tasks:** {len(ranked)} "
             f"({counts['critical']} critical, {counts['high']} high, "
             f"{counts['medium']} medium, {counts['low']} low)")
    if uncertain:
        L.append(f"- **Review manually:** {len(uncertain)} finding(s) with low confidence — see below")
    if dismissed:
        L.append(f"- **Do not fix:** {len(dismissed)} dismissed detection(s) — see the last section")
    L.append("")

    L.append("## How to use this file")
    L.append("")
    L.append("Give this file to a coding agent and tell it to work the loop. It should:")
    L.append("")
    L.append("1. Take the **first unchecked task** in order. They are sorted worst-first.")
    L.append("2. Open the file at the given path and confirm the offending code is still there. "
             "Line numbers drift — search for the evidence line rather than trusting the number.")
    L.append("3. Apply the fix described in that task, and nothing else.")
    L.append("4. Run the verification command for that task, plus the project's test suite.")
    L.append("5. Tick the checkbox, commit that single fix, and move to the next task.")
    L.append("")
    L.append("Work one task per commit. A single commit fixing thirty findings cannot be "
             "reviewed or reverted independently.")
    L.append("")

    L.append("## Rules")
    L.append("")
    L.append("- **Do not touch anything in the \"Do not fix\" section.** Those were reviewed and "
             "judged false positives. Most are deliberate test fixtures, and \"fixing\" them "
             "breaks the tests that prove the scanner works.")
    L.append("- **If the code no longer matches the evidence, stop and mark the task obsolete.** "
             "Do not hunt for something similar to change.")
    L.append("- **If a task's evidence is too thin to act on with confidence, leave it unchecked "
             "and note what you would need.** A guessed fix to security code is worse than none.")
    L.append("- **Do not widen a fix into a refactor.** Change the vulnerable construct, not the "
             "surrounding design.")
    L.append("- **Never weaken a test to make a fix pass.**")
    L.append("- Findings can be duplicates of each other. If a task is already resolved by an "
             "earlier fix, tick it and say which task covered it.")
    L.append("")

    L.append("## Progress")
    L.append("")
    for i, f in enumerate(ranked, 1):
        sev = (f.get("severity") or "unknown").upper()
        loc = f.get("file", "unknown")
        if f.get("line") not in (None, ""):
            loc += f":{f['line']}"
        cat = f.get("category") or f.get("rule") or "finding"
        L.append(f"- [ ] **{i}.** `[{sev}]` {cat} — `{loc}`")
    L.append("")

    L.append("---")
    L.append("")
    L.append("## Tasks")
    L.append("")

    for i, f in enumerate(ranked, 1):
        sev = (f.get("severity") or "unknown").upper()
        file_ = f.get("file", "unknown")
        line_no = f.get("line")
        loc = file_ + (f":{line_no}" if line_no not in (None, "") else "")
        cat = f.get("category") or f.get("rule") or "finding"
        d = details.get(index_of.get(id(f)), {}) or {}

        L.append(f"### Task {i} — [{sev}] {cat}")
        L.append("")
        L.append(f"- [ ] Fixed and verified")
        L.append("")
        L.append(f"**Where:** `{loc}`")
        L.append("")
        if f.get("description"):
            L.append(f"**Problem:** {f['description']}")
            L.append("")
        if d.get("what"):
            L.append(f"**What this is:** {d['what']}")
            L.append("")
        if d.get("why"):
            L.append(f"**Why it matters here:** {d['why']}")
            L.append("")
        if f.get("evidence"):
            L.append("**Find this line** (search for it; the line number may have moved):")
            L.append("")
            L.append("```")
            L.append(str(f["evidence"]).rstrip())
            L.append("```")
            L.append("")
        if d.get("attack"):
            L.append(f"**How it would be exploited:** {d['attack']}")
            L.append("")

        fix = d.get("fix") or f.get("improvement")
        L.append("**Change to make:**")
        L.append("")
        L.append(fix.strip() if fix else
                 "_No specific fix was supplied. Read the surrounding code, decide whether this "
                 "is real, and either fix it or leave the task unchecked with a note._")
        L.append("")
        L.append("**Verify:**")
        L.append("")
        L.append(d.get("verify")
                 or "Re-run the scan and confirm this finding is gone, then run the project's "
                    "test suite and confirm nothing regressed.")
        L.append("")
        if d.get("severity_note"):
            L.append(f"**On the severity:** {d['severity_note']}")
            L.append("")
        L.append("---")
        L.append("")

    if uncertain:
        L.append("## Review manually")
        L.append("")
        L.append("These findings have uncertain severity notes. The scanner detected something, but "
                 "the AI could not determine a confident fix. Use your judgment before fixing them.")
        L.append("")
        for i, f in enumerate(uncertain, 1):
            sev = (f.get("severity") or "unknown").upper()
            loc = f.get("file", "unknown")
            if f.get("line") not in (None, ""):
                loc += f":{f['line']}"
            cat = f.get("category") or f.get("rule") or "finding"
            d = details.get(index_of.get(id(f)), {}) or {}
            L.append(f"**{i}. [{sev}] {cat}** — `{loc}`")
            if d.get("severity_note"):
                L.append("")
                L.append(f"_Reason for review:_ {d['severity_note']}")
            L.append("")

    if dismissed:
        L.append("## Do not fix")
        L.append("")
        L.append("These were dismissed by AI review or matched inside test fixtures. They are "
                 "listed so you can recognise them if the scanner reports them again — **not** "
                 "as work. Changing them is a regression.")
        L.append("")
        by_file = {}
        for f in dismissed:
            by_file.setdefault(f.get("file") or "unknown", []).append(f)
        for fname, group in sorted(by_file.items()):
            L.append(f"**`{fname}`** — {len(group)} dismissed")
            L.append("")
            for f in group:
                ln = f.get("line")
                where = f"line {ln}" if ln not in (None, "") else "—"
                reason = f.get("ai_reason") or ("test fixture" if f.get("likely_test_fixture")
                                                else "dismissed by AI review")
                cat = f.get("category") or f.get("rule") or "finding"
                L.append(f"- {where} · {cat} — {reason}")
            L.append("")

    L.append("## When the loop is done")
    L.append("")
    L.append("1. Every task above is either ticked or has a written note saying why not.")
    L.append("2. The full test suite passes.")
    L.append("3. Re-run the scan. The new report should show the fixed findings gone, and the "
             "\"Do not fix\" list unchanged in size.")
    L.append("4. Anything still reported that you decided not to fix should be recorded, with "
             "the reason, so the next run doesn't re-litigate it.")
    L.append("")

    return "\n".join(L)


def _report_filename(label, timestamp, ext):
    """Builds a download name from what was scanned, so a folder of saved
    reports is identifiable - 'vuln-report.pdf' five times over tells you
    nothing about which scan produced which file.

    Result looks like: AI-PR-Summary_2026-08-02_1412_vuln-report.pdf
    """
    raw = (label or "").strip()
    # A local scan's label is a full path; the folder name is the useful part.
    raw = raw.replace("\\", "/").rstrip("/")
    if "/" in raw:
        raw = raw.rsplit("/", 1)[-1]
    # Keep owner/repo readable for GitHub scans that came through as owner/name.
    raw = raw.replace("(uploaded folder)", "uploaded").replace("(uploaded)", "uploaded")

    slug = re.sub(r"[^A-Za-z0-9._-]+", "-", raw).strip("-.") or "scan"
    slug = slug[:60]

    try:
        stamp = datetime.fromtimestamp(float(timestamp)).strftime("%Y-%m-%d_%H%M")
    except (TypeError, ValueError, OSError):
        stamp = datetime.now(timezone.utc).strftime("%Y-%m-%d_%H%M")

    return f"{slug}_{stamp}_vuln-report.{ext}"


@app.route("/api/vuln-scan/report", methods=["POST"])
@rate_limit("vuln_scan_report")
@require_csrf_token
def vuln_scan_report():
    """Generates a detailed vulnerability report from a completed scan: an
    AI-written narrative (executive summary, risk overview, remediation
    plan) followed by a deterministically-built, complete list of every
    finding - so nothing gets summarized away. Returned as a downloadable
    file in the requested format."""
    data = request.get_json(force=True)
    job_id = data.get("job_id", "")
    scan_ref = (data.get("scan_ref") or "").strip()
    fmt = (data.get("format") or "markdown").lower()

    # Prefer the live job (it carries the full scanner output), but fall back to
    # the persisted history so reports still work for scans from an earlier
    # session - in-memory jobs are lost on restart and pruned once over the cap.
    result = None
    # What to name the download after: the scan's custom name if it was renamed,
    # otherwise the repo or folder that was scanned.
    scan_label, scan_time = "", None
    job = VULN_JOBS.get(job_id)
    if job and job.get("status") == "done" and job.get("result"):
        result = job["result"]
        scan_label = job.get("name") or job.get("repo") or ""
        scan_time = job.get("started")
    else:
        saved = _resolve_scan(scan_ref or (f"job:{job_id}" if job_id else ""))
        if saved:
            result = {
                "findings": saved["findings"],
                "files_scanned": saved.get("files_scanned"),
                "exit_code": None,
                "from_history": True,
            }
            scan_label = saved.get("name") or saved.get("repo") or ""
            scan_time = saved.get("timestamp")

    if not result:
        saved_count = len(_load_scan_history())
        return jsonify({
            "error": (
                "No completed scan found for that reference. This scan is no longer in "
                "memory and no saved copy of it exists - scans saved before this fix "
                "didn't record their job id. "
                + (f"There are {saved_count} scan(s) in History; generate the report from "
                   "there instead (each row has pdf/docx/html/md buttons). "
                   if saved_count else "")
                + "New scans will work from this button directly."
            )
        }), 400

    findings = result.get("findings", [])
    if not findings:
        return jsonify({"error": "No findings to report on."}), 400

    client, model = get_client()
    if not client:
        return jsonify({"error": "LLM_API_KEY is not configured on the server."}), 500

    # Everything below describes findings that still need review. Dismissed
    # detections are summarised in an appendix, and must not drive the cover
    # page, the subtitle, or the AI narrative - feeding the model 139 known
    # false positives is what made it open a remediation plan with "replace
    # eval() at scan_history.json:89", a line inside its own saved scan data.
    review_findings = [f for f in findings if not _is_dismissed(f)]
    dismissed_count = len(findings) - len(review_findings)
    counts = count_findings_by_severity(review_findings)

    summary_payload = {
        "files_scanned": result.get("files_scanned"),
        "exit_code": result.get("exit_code"),
        "counts": counts,
        "total_raw_detections": len(findings),
        "dismissed_as_false_positive": dismissed_count,
        "findings": review_findings,
    }
    # The fix loop is a task list for an agent, not a document for a reader -
    # the executive narrative would only be noise in it, so don't pay for one.
    is_fix_loop = fmt in ("fixloop", "agent")
    narrative_md = ""
    if not is_fix_loop:
        findings_json = json.dumps(summary_payload, indent=2)[:REPORT_MAX_CHARS]
        try:
            narrative_md = call_llm(client, model, findings_json,
                                    system_prompt=VULN_REPORT_SYSTEM_PROMPT)
        except Exception as e:
            return jsonify({"error": f"Report generation failed: {e}"}), 500

    # Per-finding deep-dive: what it is, why it matters here, how it would be
    # exploited, the specific fix, how to verify. On by default; pass
    # detail=false for the quicker summary-only report.
    details, detail_note = {}, ""
    if data.get("detail", True):
        details, detail_note = enrich_findings_with_detail(findings)

    report_md = "" if is_fix_loop else normalize_report_text(
        narrative_md.strip() + "\n\n" + build_findings_markdown(result, details, detail_note)
    )

    title = "Vulnerability Assessment Report"
    # Name the target in the subtitle too, so an open report identifies itself
    # without needing the filename.
    target_bit = f"{scan_label} &middot; " if scan_label else ""
    dismissed_bit = f" ({dismissed_count} dismissed)" if dismissed_count else ""
    subtitle = (
        f"{target_bit}{result.get('files_scanned', 'n/a')} files scanned "
        f"&middot; {len(review_findings)} findings to review{dismissed_bit} "
        f"&middot; generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    subtitle_plain = subtitle.replace("&middot;", "-")

    def _disp(ext):
        return {"Content-Disposition":
                f'attachment; filename="{_report_filename(scan_label, scan_time, ext)}"'}

    if fmt in ("fixloop", "agent"):
        loop_md = build_fix_loop_markdown(result, details, scan_label, scan_time)
        name = _report_filename(scan_label, scan_time, "md").replace("_vuln-report.md", "_fix-loop.md")
        return Response(loop_md, mimetype="text/markdown",
                        headers={"Content-Disposition": f'attachment; filename="{name}"'})
    if fmt == "markdown":
        return Response(report_md, mimetype="text/markdown", headers=_disp("md"))
    if fmt == "html":
        return Response(build_report_html(report_md, title, subtitle), mimetype="text/html",
                         headers=_disp("html"))
    if fmt == "docx":
        return Response(build_report_docx(report_md, title, subtitle_plain),
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         headers=_disp("docx"))
    if fmt == "pdf":
        pdf_meta = {"counts": counts, "files_scanned": result.get("files_scanned"),
                    "target": scan_label}
        return Response(build_report_pdf(report_md, title, subtitle_plain, meta=pdf_meta),
                         mimetype="application/pdf", headers=_disp("pdf"))

    return jsonify({"error": f"Unknown format: {fmt}"}), 400


if __name__ == "__main__":
    # debug=True enables the Werkzeug interactive debugger, which hands out a
    # Python console on any unhandled exception - that is arbitrary code
    # execution for anyone who can reach the port. It also made the reloader
    # restart mid-scan whenever a file changed, losing running jobs. Off by
    # default now; set DASHBOARD_DEBUG=1 when you actually want it.
    debug = os.environ.get("DASHBOARD_DEBUG", "0") == "1"
    if debug:
        print("WARNING: debug mode is on - the Werkzeug debugger allows code "
              "execution. Never use this on a shared or reachable machine.")
    # host stays on loopback so the API (which reads local folders and runs
    # build commands) is never exposed to the local network.
    app.run(
        host="127.0.0.1", port=5000, debug=debug,
        # exclude_patterns keeps the dev reloader from watching .repo-cache/ -
        # without it, checking out a branch in a cached clone (which touches
        # file mtimes, including any server.py inside a cloned repo) can
        # trigger an unwanted restart mid-request.
        exclude_patterns=["*/.repo-cache/*", "*\\.repo-cache\\*"],
    )
